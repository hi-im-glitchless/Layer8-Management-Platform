"""DOCX instruction applier -- modifies DOCX templates in-place preserving formatting.

Receives structured JSON instructions (from the LLM via the rules engine) and
applies Jinja2 placeholders to a client DOCX template using python-docx while
preserving all original formatting (fonts, sizes, colors, alignment).
"""
import copy
import logging
from io import BytesIO

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph as _ParagraphWrapper
from lxml import etree

from app.models.adapter import Instruction, InstructionSet

logger = logging.getLogger(__name__)

# OOXML namespace constants
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# Fields that repeat across headers, footers, and text boxes.
# Only these get the replace-all-occurrences treatment (strategies 3-6).
_REPEATING_FIELDS: set[str] = {
    "report_date", "client.short_name", "title", "project.codename",
    "project.start_date", "project.end_date",
    "team[0].name", "team[0].email",
}


class InstructionApplier:
    """Applies structured instructions to modify a DOCX template.

    Handles four instruction actions:
      - replace_text: Replace text within a paragraph, preserving run formatting
      - insert_before: Insert a new paragraph before the target
      - insert_after: Insert a new paragraph after the target
      - wrap_table_row: Add loop markers around a table row via XML manipulation
    """

    def apply(
        self,
        template_bytes: bytes,
        instruction_set: InstructionSet,
    ) -> tuple[bytes, int, int, list[str]]:
        """Apply instructions to a DOCX template.

        Instructions are sorted by paragraph_index descending and applied
        from bottom to top to avoid index shifts.

        Args:
            template_bytes: Raw bytes of the client DOCX template.
            instruction_set: The validated and enriched instruction set.

        Returns:
            Tuple of (modified_docx_bytes, applied_count, skipped_count, warnings).
        """
        doc = Document(BytesIO(template_bytes))
        body_paragraphs = doc.paragraphs

        applied = 0
        skipped = 0
        warnings: list[str] = []

        # Sort instructions by paragraph_index descending to avoid index shifts
        sorted_instructions = sorted(
            instruction_set.instructions,
            key=lambda inst: inst.paragraph_index,
            reverse=True,
        )

        for inst in sorted_instructions:
            try:
                success = self._apply_instruction(doc, body_paragraphs, inst)
                if success:
                    applied += 1
                else:
                    skipped += 1
                    warnings.append(
                        f"Skipped instruction at paragraph {inst.paragraph_index}: "
                        f"action={inst.action}, could not apply"
                    )
            except Exception as exc:
                skipped += 1
                warnings.append(
                    f"Error applying instruction at paragraph {inst.paragraph_index}: "
                    f"{exc}"
                )
                logger.warning(
                    "Failed to apply instruction at paragraph %d: %s",
                    inst.paragraph_index,
                    exc,
                )

        # Save to bytes
        output = BytesIO()
        doc.save(output)
        return output.getvalue(), applied, skipped, warnings

    def _apply_instruction(
        self,
        doc: Document,
        body_paragraphs: list,
        instruction: Instruction,
    ) -> bool:
        """Apply a single instruction to the document.

        For replace_text actions, uses a multi-strategy approach:
          1. Try the body paragraph at the given index
          2. Search all body paragraphs by text
          3. Search headers and footers by text

        Returns True if the instruction was applied successfully, False otherwise.
        """
        idx = instruction.paragraph_index

        if instruction.action == "replace_text":
            return self._replace_text_with_fallback(
                doc, body_paragraphs, instruction,
            )

        # For non-replace actions, use body paragraph at the given index
        if idx < 0 or idx >= len(body_paragraphs):
            logger.warning(
                "Paragraph index %d out of range (0-%d)",
                idx,
                len(body_paragraphs) - 1,
            )
            return False

        paragraph = body_paragraphs[idx]

        if instruction.action == "insert_before":
            return self._insert_paragraph(
                doc, paragraph, instruction.replacement_text, before=True
            )
        elif instruction.action == "insert_after":
            return self._insert_paragraph(
                doc, paragraph, instruction.replacement_text, before=False
            )
        elif instruction.action == "wrap_table_row":
            return self._wrap_table_row_with_loop(
                doc, paragraph, instruction.replacement_text
            )
        else:
            logger.warning("Unknown action: %s", instruction.action)
            return False

    def _replace_text_with_fallback(
        self,
        doc: Document,
        body_paragraphs: list,
        instruction: Instruction,
    ) -> bool:
        """Replace text using multi-strategy search across the document.

        All strategies replace every occurrence found.  Header/footer
        scanning (strategies 3 and 6) is restricted to repeating fields
        since content fields never appear there.

        Strategy order:
        1. Body paragraph at the given index
        2. All other body paragraphs
        3. Header/footer top-level paragraphs (repeating only)
        4. Body table cell paragraphs
        5. Body text box paragraphs (w:txbxContent)
        6. Header/footer nested tables and text boxes (repeating only)
        """
        idx = instruction.paragraph_index
        original = instruction.original_text
        replacement = instruction.replacement_text
        is_repeating = instruction.gw_field in _REPEATING_FIELDS
        total_replaced = 0

        # Strategy 1: Try body paragraph at the given index
        if 0 <= idx < len(body_paragraphs):
            if self._replace_in_paragraph(body_paragraphs[idx], original, replacement):
                total_replaced += 1

        # Strategy 2: Search all other body paragraphs
        for i, para in enumerate(body_paragraphs):
            if i == idx:
                continue
            if self._replace_in_paragraph(para, original, replacement):
                logger.info(
                    "Replaced occurrence in paragraph %d (gw_field=%s)",
                    i, instruction.gw_field,
                )
                total_replaced += 1

        # Strategy 3: Header/footer top-level paragraphs (repeating only)
        if is_repeating:
            for section in doc.sections:
                for hf_label, hf in self._iter_header_footers(section):
                    for para in hf.paragraphs:
                        if self._replace_in_paragraph(para, original, replacement):
                            logger.info(
                                "Applied replacement in %s (gw_field=%s)",
                                hf_label, instruction.gw_field,
                            )
                            total_replaced += 1

        # Strategy 4: Body table cell paragraphs
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if self._replace_in_paragraph(para, original, replacement):
                            logger.info(
                                "Applied replacement in table cell (gw_field=%s)",
                                instruction.gw_field,
                            )
                            total_replaced += 1

        # Strategy 5: Body text boxes
        for txbx in doc.element.body.findall(".//" + qn("w:txbxContent")):
            for p_elem in txbx.findall(qn("w:p")):
                para = _ParagraphWrapper(p_elem, doc.element.body)
                if self._replace_in_paragraph(para, original, replacement):
                    logger.info(
                        "Applied replacement in text box (gw_field=%s)",
                        instruction.gw_field,
                    )
                    total_replaced += 1

        # Strategy 6: Header/footer nested tables and text boxes (repeating only)
        if is_repeating:
            for section in doc.sections:
                for hf_label, hf in self._iter_header_footers(section):
                    hf_elem = hf._element
                    for tc_elem in hf_elem.findall(".//" + qn("w:tc")):
                        for p_elem in tc_elem.findall(qn("w:p")):
                            para = _ParagraphWrapper(p_elem, hf_elem)
                            if self._replace_in_paragraph(para, original, replacement):
                                logger.info(
                                    "Applied replacement in %s table (gw_field=%s)",
                                    hf_label, instruction.gw_field,
                                )
                                total_replaced += 1
                    for txbx in hf_elem.findall(".//" + qn("w:txbxContent")):
                        for p_elem in txbx.findall(qn("w:p")):
                            para = _ParagraphWrapper(p_elem, hf_elem)
                            if self._replace_in_paragraph(para, original, replacement):
                                logger.info(
                                    "Applied replacement in %s text box (gw_field=%s)",
                                    hf_label, instruction.gw_field,
                                )
                                total_replaced += 1

        if total_replaced > 1:
            logger.info(
                "Replaced %d occurrences of '%s' -> '%s' (gw_field=%s)",
                total_replaced, original[:50], replacement[:50], instruction.gw_field,
            )

        return total_replaced > 0

    @staticmethod
    def _iter_header_footers(section):
        """Yield (label, header_or_footer) for all variants of a section.

        Covers default, first-page, and even-page headers and footers.
        """
        for label, accessor in [
            ("header", "header"),
            ("footer", "footer"),
            ("first-page header", "first_page_header"),
            ("first-page footer", "first_page_footer"),
            ("even-page header", "even_page_header"),
            ("even-page footer", "even_page_footer"),
        ]:
            try:
                hf = getattr(section, accessor)
                if hf is not None:
                    yield label, hf
            except Exception:
                continue

    # ------------------------------------------------------------------
    # replace_text
    # ------------------------------------------------------------------

    def _replace_in_paragraph(
        self,
        paragraph,
        original_text: str,
        replacement_text: str,
    ) -> bool:
        """Replace original_text with replacement_text in a paragraph, preserving formatting.

        Handles text that may be split across multiple runs (common in Word-generated DOCX).

        Strategy:
          1. If original_text is entirely within one run, replace in that run.
          2. If original_text spans multiple runs, consolidate into the first
             matching run and remove consumed text from subsequent runs.

        Returns True if replacement was made, False otherwise.
        """
        runs = paragraph.runs
        if not runs:
            return False

        # Strategy 1: Check if original_text is within a single run
        for run in runs:
            if original_text in run.text:
                run.text = run.text.replace(original_text, replacement_text, 1)
                return True

        # Strategy 2: Text spans multiple runs -- find the span
        full_text = "".join(r.text for r in runs)
        start_pos = full_text.find(original_text)
        if start_pos == -1:
            return False

        end_pos = start_pos + len(original_text)

        # Map character positions to runs
        run_boundaries = self._build_run_boundaries(runs)

        # Find which runs are involved
        first_run_idx = None
        last_run_idx = None
        for i, (run_start, run_end) in enumerate(run_boundaries):
            if first_run_idx is None and run_end > start_pos:
                first_run_idx = i
            if run_end >= end_pos:
                last_run_idx = i
                break

        if first_run_idx is None or last_run_idx is None:
            return False

        # Build the replacement by modifying the first run and cleaning up others
        first_run = runs[first_run_idx]
        first_run_start, first_run_end = run_boundaries[first_run_idx]

        # Text before the match in the first run
        prefix = first_run.text[: start_pos - first_run_start]
        # Text after the match in the last run
        last_run_start, last_run_end = run_boundaries[last_run_idx]
        suffix = runs[last_run_idx].text[end_pos - last_run_start:]

        # Set the first run's text to prefix + replacement + suffix (if same run)
        if first_run_idx == last_run_idx:
            first_run.text = prefix + replacement_text + suffix
        else:
            first_run.text = prefix + replacement_text
            # Clear intermediate runs
            for i in range(first_run_idx + 1, last_run_idx):
                runs[i].text = ""
            # Set the last run to just the suffix
            runs[last_run_idx].text = suffix

        return True

    @staticmethod
    def _build_run_boundaries(runs: list) -> list[tuple[int, int]]:
        """Build a list of (start, end) character positions for each run."""
        boundaries = []
        pos = 0
        for run in runs:
            length = len(run.text)
            boundaries.append((pos, pos + length))
            pos += length
        return boundaries

    # ------------------------------------------------------------------
    # insert_before / insert_after
    # ------------------------------------------------------------------

    def _insert_paragraph(
        self,
        doc: Document,
        target_paragraph,
        text: str,
        before: bool = True,
    ) -> bool:
        """Insert a new paragraph before or after the target paragraph.

        Copies the target paragraph's style to the new paragraph.

        Args:
            doc: The Document object.
            target_paragraph: The reference paragraph.
            text: The text content for the new paragraph.
            before: If True, insert before; if False, insert after.

        Returns:
            True if insertion succeeded.
        """
        target_elem = target_paragraph._element
        parent = target_elem.getparent()

        if parent is None:
            return False

        # Create a new paragraph element
        new_p = copy.deepcopy(target_elem)
        # Clear runs from the copied paragraph
        for r in new_p.findall(qn("w:r")):
            new_p.remove(r)

        # Add a single run with the replacement text
        new_run = etree.SubElement(new_p, qn("w:r"))

        # Copy run properties from the first run of the target if available
        target_runs = target_elem.findall(qn("w:r"))
        if target_runs:
            source_rpr = target_runs[0].find(qn("w:rPr"))
            if source_rpr is not None:
                new_run.insert(0, copy.deepcopy(source_rpr))

        new_t = etree.SubElement(new_run, qn("w:t"))
        new_t.text = text
        new_t.set(qn("xml:space"), "preserve")

        if before:
            parent.insert(list(parent).index(target_elem), new_p)
        else:
            parent.insert(list(parent).index(target_elem) + 1, new_p)

        return True

    # ------------------------------------------------------------------
    # wrap_table_row
    # ------------------------------------------------------------------

    def _wrap_table_row_with_loop(
        self,
        doc: Document,
        paragraph,
        loop_expr: str,
    ) -> bool:
        """Wrap the table row containing the target paragraph with loop markers.

        Navigates from the paragraph to its parent table row via XML tree,
        then inserts loop start/end marker rows.

        Args:
            doc: The Document object.
            paragraph: The paragraph inside the table cell.
            loop_expr: The loop expression (e.g., "{%tr for item in scope %}...{%tr endfor %}").

        Returns:
            True if wrapping succeeded, False if paragraph is not in a table.
        """
        para_elem = paragraph._element

        # Navigate up to find the table row (w:tr) containing this paragraph
        tr_elem = self._find_ancestor(para_elem, qn("w:tr"))
        if tr_elem is None:
            logger.warning("Paragraph is not inside a table row; cannot wrap")
            return False

        tbl_elem = self._find_ancestor(tr_elem, qn("w:tbl"))
        if tbl_elem is None:
            logger.warning("Table row has no parent table element")
            return False

        # Parse the loop expression to extract start and end markers
        start_marker, end_marker = self._parse_loop_markers(loop_expr)

        # Create marker rows -- minimal rows with a single cell containing the marker text
        tr_index = list(tbl_elem).index(tr_elem)

        start_row = self._create_marker_row(tr_elem, start_marker)
        end_row = self._create_marker_row(tr_elem, end_marker)

        # Insert: start_row before the data row, end_row after
        tbl_elem.insert(tr_index, start_row)
        # After inserting start_row, the data row shifted by 1
        tbl_elem.insert(tr_index + 2, end_row)

        return True

    @staticmethod
    def _find_ancestor(elem, tag: str):
        """Walk up the XML tree to find an ancestor with the given tag."""
        parent = elem.getparent()
        while parent is not None:
            if parent.tag == tag:
                return parent
            parent = parent.getparent()
        return None

    @staticmethod
    def _parse_loop_markers(loop_expr: str) -> tuple[str, str]:
        """Parse a loop expression into start and end markers.

        Handles formats like:
          - "{%tr for item in scope %}" -> ("{%tr for item in scope %}", "{%tr endfor %}")
          - "for item in scope" -> ("{%tr for item in scope %}", "{%tr endfor %}")
        """
        stripped = loop_expr.strip()

        # If it contains both start and end markers, split them
        if "{%tr endfor" in stripped or "{% endfor" in stripped:
            # Split on endfor marker
            parts = stripped.split("{%tr endfor")
            if len(parts) == 2:
                return parts[0].strip(), "{%tr endfor %}"
            parts = stripped.split("{% endfor")
            if len(parts) == 2:
                return parts[0].strip(), "{% endfor %}"

        # If it's just the start expression
        if stripped.startswith("{%tr"):
            return stripped, "{%tr endfor %}"
        elif stripped.startswith("{%"):
            # Convert to {%tr format
            inner = stripped[2:].rstrip("%}").strip()
            return "{%tr " + inner + " %}", "{%tr endfor %}"
        else:
            # Raw expression like "for item in scope"
            return "{%tr " + stripped + " %}", "{%tr endfor %}"

    @staticmethod
    def _create_marker_row(template_row, marker_text: str):
        """Create a table row element containing a single cell with marker text.

        The marker row copies the structure (number of cells) from the template row
        but only puts text in the first cell.
        """
        new_tr = copy.deepcopy(template_row)

        # Clear all cell content
        cells = new_tr.findall(qn("w:tc"))
        for i, tc in enumerate(cells):
            # Remove all paragraphs
            for p in tc.findall(qn("w:p")):
                tc.remove(p)

            # Add a new empty paragraph (or with marker in first cell)
            new_p = etree.SubElement(tc, qn("w:p"))
            if i == 0:
                new_r = etree.SubElement(new_p, qn("w:r"))
                new_t = etree.SubElement(new_r, qn("w:t"))
                new_t.text = marker_text
                new_t.set(qn("xml:space"), "preserve")

        return new_tr

    # ------------------------------------------------------------------
    # Formatting preservation utilities
    # ------------------------------------------------------------------

    @staticmethod
    def _preserve_formatting(source_run, target_run) -> None:
        """Copy font properties from source_run to target_run.

        Copies: name, size, bold, italic, underline, color, strikethrough.

        Args:
            source_run: python-docx Run to copy formatting from.
            target_run: python-docx Run to apply formatting to.
        """
        source_font = source_run.font
        target_font = target_run.font

        # Copy font properties
        if source_font.name is not None:
            target_font.name = source_font.name
        if source_font.size is not None:
            target_font.size = source_font.size
        if source_font.bold is not None:
            target_font.bold = source_font.bold
        if source_font.italic is not None:
            target_font.italic = source_font.italic
        if source_font.underline is not None:
            target_font.underline = source_font.underline
        if source_font.color and source_font.color.rgb is not None:
            target_font.color.rgb = source_font.color.rgb
        if source_font.strike is not None:
            target_font.strike = source_font.strike
