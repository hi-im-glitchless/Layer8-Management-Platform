"""Annotated DOCX preview service -- applies paragraph shading and generates metadata.

Applies green (mapped) and yellow (gap) background shading to DOCX paragraphs
via OOXML XML manipulation, and generates tooltip/unmapped-paragraph metadata
for the frontend overlay.

Also provides placeholder preview shading (light blue) for adapted DOCX files
containing Jinja2 expressions after auto-map insertion.
"""
import logging
import re
from io import BytesIO

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

from app.models.adapter import MappingPlan, PlaceholderInfo
from app.models.gap_detection import (
    AnnotationMetadata,
    GapEntry,
    TooltipEntry,
    UnmappedParagraph,
)
from app.services.docx_parser import DocxParserService

logger = logging.getLogger(__name__)

# Shading colors (hex, no leading #)
GREEN_SHADING = "C6EFCE"       # RGB(198, 239, 206) -- mapped paragraphs
YELLOW_SHADING = "FFF2CC"      # RGB(255, 242, 204) -- gap paragraphs
PLACEHOLDER_SHADING = "DAEAF6" # RGB(218, 234, 246) -- placeholder paragraphs (light blue)

# OOXML namespace
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def apply_paragraph_shading(
    doc_bytes: bytes,
    mapping_plan: MappingPlan,
    gaps: list[GapEntry],
    green_only: bool = False,
) -> bytes:
    """Apply green/yellow paragraph background shading to a DOCX document.

    Args:
        doc_bytes: Raw bytes of the client DOCX template.
        mapping_plan: The validated mapping plan with section_index references.
        gaps: List of detected gaps with estimated paragraph indices.
        green_only: When True, skip yellow gap shading and only apply green
            to mapped paragraphs. Used in interactive mapping mode.

    Returns:
        Modified DOCX bytes with paragraph shading applied.
    """
    doc = Document(BytesIO(doc_bytes))
    paragraphs = doc.paragraphs

    # Apply green shading to mapped paragraphs
    mapped_indices: set[int] = set()
    for entry in mapping_plan.entries:
        idx = entry.section_index
        if 0 <= idx < len(paragraphs):
            _set_paragraph_shading(paragraphs[idx], GREEN_SHADING)
            mapped_indices.add(idx)
        else:
            logger.warning(
                "Mapped section_index %d out of range (0-%d), skipping shading",
                idx,
                len(paragraphs) - 1,
            )

    # Apply yellow shading to gap paragraphs (skip when green_only=True)
    if not green_only:
        for gap in gaps:
            idx = gap.estimated_paragraph_index
            if idx is not None and 0 <= idx < len(paragraphs) and idx not in mapped_indices:
                _set_paragraph_shading(paragraphs[idx], YELLOW_SHADING)
            elif idx is not None:
                logger.debug(
                    "Gap paragraph index %d out of range or already mapped, skipping",
                    idx,
                )

    # Save modified document to bytes
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def _set_paragraph_shading(paragraph, hex_color: str) -> None:
    """Apply background shading to a paragraph via OOXML XML manipulation.

    Adds or updates the w:shd element within the paragraph's w:pPr element:
    <w:shd w:val="clear" w:color="auto" w:fill="{hex_color}"/>

    Args:
        paragraph: A python-docx Paragraph object.
        hex_color: Six-character hex color string (e.g. "C6EFCE").
    """
    p_elem = paragraph._element

    # Get or create w:pPr (paragraph properties)
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = etree.SubElement(p_elem, qn("w:pPr"))
        # Insert pPr as first child (OOXML schema requires it before runs)
        p_elem.insert(0, pPr)

    # Get or create w:shd (shading) within pPr
    shd = pPr.find(qn("w:shd"))
    if shd is None:
        shd = etree.SubElement(pPr, qn("w:shd"))

    # Set shading attributes
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)


def generate_annotation_metadata(
    doc_bytes: bytes,
    mapping_plan: MappingPlan,
    gaps: list[GapEntry],
) -> AnnotationMetadata:
    """Generate tooltip data and unmapped paragraph list for the frontend.

    Args:
        doc_bytes: Raw bytes of the client DOCX template.
        mapping_plan: The validated mapping plan.
        gaps: List of detected gaps from gap detection.

    Returns:
        AnnotationMetadata with tooltip_data and unmapped_paragraphs.
    """
    parser = DocxParserService()
    doc_structure = parser.parse(doc_bytes)
    paragraphs = doc_structure.paragraphs

    # Track which paragraph indices are accounted for (mapped or gap)
    accounted_indices: set[int] = set()

    # Build tooltip entries for mapped paragraphs
    tooltip_data: list[TooltipEntry] = []
    for entry in mapping_plan.entries:
        idx = entry.section_index
        if 0 <= idx < len(paragraphs):
            tooltip_data.append(
                TooltipEntry(
                    paragraph_index=idx,
                    gw_field=entry.gw_field,
                    marker_type=entry.marker_type,
                    section_text=entry.section_text,
                    status="mapped",
                )
            )
            accounted_indices.add(idx)

    # Build tooltip entries for gap paragraphs
    for gap in gaps:
        idx = gap.estimated_paragraph_index
        if idx is not None and 0 <= idx < len(paragraphs) and idx not in accounted_indices:
            tooltip_data.append(
                TooltipEntry(
                    paragraph_index=idx,
                    gw_field=gap.gw_field,
                    marker_type=gap.marker_type,
                    section_text=paragraphs[idx].text[:200],
                    status="gap",
                )
            )
            accounted_indices.add(idx)

    # Sort tooltip data by paragraph index
    tooltip_data.sort(key=lambda t: t.paragraph_index)

    # Build unmapped paragraphs list (not in mapping plan, not a gap, not empty)
    unmapped: list[UnmappedParagraph] = []
    for i, para in enumerate(paragraphs):
        if i in accounted_indices:
            continue
        text = para.text.strip()
        if not text:
            continue
        unmapped.append(
            UnmappedParagraph(
                paragraph_index=i,
                text=text[:200],
                heading_level=para.heading_level,
            )
        )

    # Already sorted by paragraph_index (iterated in order)

    return AnnotationMetadata(
        tooltip_data=tooltip_data,
        unmapped_paragraphs=unmapped,
    )


# ---------------------------------------------------------------------------
# Placeholder Preview (Phase 5.3)
# ---------------------------------------------------------------------------

# Regex to match Jinja2 expressions: {{ ... }}
_JINJA2_PATTERN = re.compile(r"\{\{.*?\}\}")


def generate_placeholder_preview(
    doc_bytes: bytes,
) -> tuple[bytes, list[PlaceholderInfo]]:
    """Generate a placeholder-styled preview of an adapted DOCX.

    Scans all paragraphs (including those inside tables) for Jinja2
    expressions (``{{ ... }}``), applies light blue background shading
    to paragraphs that contain at least one placeholder, and returns
    the annotated DOCX bytes plus a list of placeholder metadata.

    Args:
        doc_bytes: Raw bytes of the adapted DOCX with Jinja2 placeholders.

    Returns:
        Tuple of (annotated_docx_bytes, placeholder_info_list).
    """
    doc = Document(BytesIO(doc_bytes))
    placeholders: list[PlaceholderInfo] = []

    # Scan body paragraphs
    for idx, paragraph in enumerate(doc.paragraphs):
        _scan_paragraph_for_placeholders(paragraph, idx, placeholders)

    # Scan table cell paragraphs
    # Use a running index that continues after body paragraphs
    table_para_offset = len(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _scan_paragraph_for_placeholders(
                        paragraph, table_para_offset, placeholders
                    )
                    table_para_offset += 1

    # Save modified document to bytes
    output = BytesIO()
    doc.save(output)
    annotated_bytes = output.getvalue()

    logger.info(
        "Placeholder preview: %d placeholders found across %d paragraphs",
        len(placeholders),
        len({p.paragraph_index for p in placeholders}),
    )

    return annotated_bytes, placeholders


def _set_run_shading(run, hex_color: str) -> None:
    """Apply background shading to a single run via OOXML XML manipulation.

    Adds or updates the w:shd element within the run's w:rPr element:
    <w:shd w:val="clear" w:color="auto" w:fill="{hex_color}"/>

    Args:
        run: A python-docx Run object.
        hex_color: Six-character hex color string (e.g. "DAEAF6").
    """
    r_elem = run._element

    # Get or create w:rPr (run properties)
    rPr = r_elem.find(qn("w:rPr"))
    if rPr is None:
        rPr = etree.SubElement(r_elem, qn("w:rPr"))
        # Insert rPr as first child (OOXML schema requires it before w:t)
        r_elem.insert(0, rPr)

    # Get or create w:shd (shading) within rPr
    shd = rPr.find(qn("w:shd"))
    if shd is None:
        shd = etree.SubElement(rPr, qn("w:shd"))

    # Set shading attributes
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)


def _scan_paragraph_for_placeholders(
    paragraph,
    paragraph_index: int,
    placeholders: list[PlaceholderInfo],
) -> None:
    """Scan a single paragraph for Jinja2 expressions and apply shading.

    If the paragraph text contains one or more ``{{ ... }}`` patterns,
    light blue shading is applied to the specific runs containing them
    (not the entire paragraph) and a PlaceholderInfo entry is added
    for each match.

    Args:
        paragraph: A python-docx Paragraph object.
        paragraph_index: The index to record in PlaceholderInfo.
        placeholders: Accumulator list -- matched entries are appended.
    """
    text = paragraph.text
    matches = _JINJA2_PATTERN.findall(text)
    if not matches:
        return

    # Apply light blue shading only to runs that contain Jinja2 expressions
    for run in paragraph.runs:
        if _JINJA2_PATTERN.search(run.text):
            _set_run_shading(run, PLACEHOLDER_SHADING)

    # Extract placeholder info for each match
    for match in matches:
        # Strip {{ }} and whitespace to get the field path
        gw_field = match.strip().removeprefix("{{").removesuffix("}}").strip()
        placeholders.append(
            PlaceholderInfo(
                paragraph_index=paragraph_index,
                placeholder_text=match,
                gw_field=gw_field,
            )
        )
