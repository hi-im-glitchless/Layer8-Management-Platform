"""DOCX parsing service -- extracts structured content from DOCX files."""
import logging
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.shared import Emu

from app.models.docx import (
    DocxCell,
    DocxImage,
    DocxParagraph,
    DocxRow,
    DocxRun,
    DocxSection,
    DocxStructure,
    DocxTable,
)

logger = logging.getLogger(__name__)


class DocxParserService:
    """Parses DOCX files into structured JSON representations."""

    def parse(self, file_bytes: bytes) -> DocxStructure:
        """Parse a DOCX file from bytes and return structured content.

        Args:
            file_bytes: Raw bytes of the .docx file.

        Returns:
            DocxStructure with paragraphs, tables, images, sections, styles, metadata.

        Raises:
            ValueError: If the bytes cannot be parsed as a valid DOCX.
        """
        try:
            doc = Document(BytesIO(file_bytes))
        except Exception as exc:
            raise ValueError(f"Failed to parse DOCX: {exc}") from exc

        paragraphs = self._extract_paragraphs(doc)
        tables = self._extract_tables(doc)
        images = self._extract_images(doc)
        sections = self._extract_sections(doc)
        styles = self._extract_styles(doc)
        metadata = self._extract_metadata(doc)

        return DocxStructure(
            paragraphs=paragraphs,
            tables=tables,
            images=images,
            sections=sections,
            styles=styles,
            metadata=metadata,
        )

    # ------------------------------------------------------------------
    # Paragraph extraction
    # ------------------------------------------------------------------

    def _extract_paragraphs(self, doc: Document) -> list[DocxParagraph]:
        """Extract all body paragraphs with text, style, heading level, runs, and zone tags.

        Cover detection: paragraphs before the first heading (or before index 15,
        whichever comes first) get zone="cover" instead of "body".
        """
        result: list[DocxParagraph] = []

        # First pass: find the index of the first heading for cover detection
        first_heading_idx: int | None = None
        for idx, para in enumerate(doc.paragraphs):
            if idx >= 15:
                break
            level = self._heading_level(para)
            if level is not None:
                first_heading_idx = idx
                break

        for idx, para in enumerate(doc.paragraphs):
            parsed = self._parse_paragraph(para)
            # Determine zone: cover for pre-heading paragraphs, body otherwise
            if first_heading_idx is not None and idx < first_heading_idx:
                parsed.zone = "cover"
            else:
                parsed.zone = "body"
            result.append(parsed)

        return result

    def _parse_paragraph(self, para: Any) -> DocxParagraph:
        """Convert a python-docx Paragraph object to a DocxParagraph model."""
        heading_level = self._heading_level(para)
        alignment = para.alignment.name if para.alignment is not None else None
        runs = [self._parse_run(r) for r in para.runs]

        return DocxParagraph(
            text=para.text,
            style_name=para.style.name if para.style else None,
            heading_level=heading_level,
            alignment=alignment,
            runs=runs,
        )

    @staticmethod
    def _heading_level(para: Any) -> int | None:
        """Derive heading level from paragraph style name (e.g. 'Heading 2' -> 2)."""
        style = para.style
        if style is None:
            return None
        name = style.name or ""
        if name.startswith("Heading"):
            parts = name.split()
            if len(parts) == 2 and parts[1].isdigit():
                return int(parts[1])
        return None

    @staticmethod
    def _parse_run(run: Any) -> DocxRun:
        """Convert a python-docx Run object to a DocxRun model."""
        font = run.font
        color_hex: str | None = None
        if font.color and font.color.rgb:
            color_hex = str(font.color.rgb)

        font_size: float | None = None
        if font.size is not None:
            font_size = font.size.pt

        return DocxRun(
            text=run.text,
            bold=font.bold,
            italic=font.italic,
            underline=font.underline,
            font_name=font.name,
            font_size=font_size,
            color=color_hex,
        )

    # ------------------------------------------------------------------
    # Table extraction
    # ------------------------------------------------------------------

    def _extract_tables(self, doc: Document) -> list[DocxTable]:
        """Extract all tables with rows, cells, and nested paragraphs.

        Cell paragraphs are tagged with zone="table_cell" and table_index
        set to the table's position in the document.
        """
        result: list[DocxTable] = []
        for tbl_idx, table in enumerate(doc.tables):
            rows: list[DocxRow] = []
            for row in table.rows:
                cells: list[DocxCell] = []
                for cell in row.cells:
                    cell_paras: list[DocxParagraph] = []
                    for p in cell.paragraphs:
                        parsed = self._parse_paragraph(p)
                        parsed.zone = "table_cell"
                        parsed.table_index = tbl_idx
                        cell_paras.append(parsed)
                    merge_info = self._get_merge_info(cell)
                    cells.append(
                        DocxCell(
                            text=cell.text,
                            paragraphs=cell_paras,
                            merge_info=merge_info,
                        )
                    )
                rows.append(DocxRow(cells=cells))

            style_name = table.style.name if table.style else None
            result.append(DocxTable(rows=rows, style_name=style_name))
        return result

    @staticmethod
    def _get_merge_info(cell: Any) -> dict | None:
        """Inspect cell XML for vertical/horizontal merge attributes."""
        tc = cell._tc
        merge: dict[str, str] = {}

        v_merge = tc.find(qn("w:tcPr"))
        if v_merge is not None:
            vm_elem = v_merge.find(qn("w:vMerge"))
            if vm_elem is not None:
                merge["vertical"] = vm_elem.get(qn("w:val"), "continue")

            grid_span = v_merge.find(qn("w:gridSpan"))
            if grid_span is not None:
                val = grid_span.get(qn("w:val"))
                if val and int(val) > 1:
                    merge["horizontal_span"] = val

        return merge if merge else None

    # ------------------------------------------------------------------
    # Image extraction
    # ------------------------------------------------------------------

    def _extract_images(self, doc: Document) -> list[DocxImage]:
        """Extract inline image references with content type and dimensions."""
        images: list[DocxImage] = []

        for para_idx, para in enumerate(doc.paragraphs):
            for run in para.runs:
                inline_shapes = run._element.findall(
                    f".//{qn('wp:inline')}"
                )
                for inline in inline_shapes:
                    img = self._parse_inline_image(inline, para_idx, doc)
                    if img is not None:
                        images.append(img)

        return images

    def _parse_inline_image(
        self, inline: Any, para_idx: int, doc: Document
    ) -> DocxImage | None:
        """Parse an inline element for image data."""
        try:
            extent = inline.find(qn("wp:extent"))
            width = int(extent.get("cx", 0)) if extent is not None else None
            height = int(extent.get("cy", 0)) if extent is not None else None

            # Navigate to the blip element for the image relationship ID
            graphic = inline.find(f".//{qn('a:blip')}")
            if graphic is None:
                return None

            r_embed = graphic.get(qn("r:embed"))
            if r_embed is None:
                return None

            # Resolve relationship to get content type and filename
            part = doc.part
            rel = part.rels.get(r_embed)
            if rel is None:
                return None

            target_part = rel.target_part
            content_type = target_part.content_type if target_part else None
            filename = target_part.partname.split("/")[-1] if target_part else None

            return DocxImage(
                content_type=content_type,
                width=width,
                height=height,
                filename=filename,
                paragraph_index=para_idx,
            )
        except Exception:
            logger.debug("Failed to parse inline image at paragraph %d", para_idx)
            return None

    # ------------------------------------------------------------------
    # Section extraction
    # ------------------------------------------------------------------

    def _extract_sections(self, doc: Document) -> list[DocxSection]:
        """Extract section properties including headers, footers, and page layout."""
        result: list[DocxSection] = []
        for section in doc.sections:
            header_paras = self._extract_header_footer_paras(section, "header")
            footer_paras = self._extract_header_footer_paras(section, "footer")

            orientation = None
            if section.orientation is not None:
                orientation = (
                    "LANDSCAPE"
                    if section.orientation == WD_ORIENT.LANDSCAPE
                    else "PORTRAIT"
                )

            page_width = section.page_width
            page_height = section.page_height

            result.append(
                DocxSection(
                    header_paragraphs=header_paras,
                    footer_paragraphs=footer_paras,
                    page_width=page_width,
                    page_height=page_height,
                    orientation=orientation,
                )
            )
        return result

    def _extract_header_footer_paras(
        self, section: Any, part_type: str
    ) -> list[DocxParagraph]:
        """Extract paragraphs from a section's header or footer.

        Tags each paragraph with zone="header" or zone="footer".
        """
        zone = "header" if part_type == "header" else "footer"
        paras: list[DocxParagraph] = []
        try:
            hf = section.header if part_type == "header" else section.footer
            if hf is None or not hf.is_linked_to_previous:
                pass  # Extract even linked headers for completeness
            if hf is not None:
                for para in hf.paragraphs:
                    parsed = self._parse_paragraph(para)
                    parsed.zone = zone
                    paras.append(parsed)
        except Exception:
            logger.debug("Failed to extract %s paragraphs", part_type)
        return paras

    # ------------------------------------------------------------------
    # Style and metadata extraction
    # ------------------------------------------------------------------

    @staticmethod
    def _extract_styles(doc: Document) -> list[str]:
        """Collect unique style names used across the document."""
        style_names: set[str] = set()
        for para in doc.paragraphs:
            if para.style and para.style.name:
                style_names.add(para.style.name)
        for table in doc.tables:
            if table.style and table.style.name:
                style_names.add(table.style.name)
        return sorted(style_names)

    @staticmethod
    def _extract_metadata(doc: Document) -> dict:
        """Read document core properties as a plain dict."""
        meta: dict[str, Any] = {}
        props = doc.core_properties
        if props is None:
            return meta

        for attr in (
            "author",
            "title",
            "subject",
            "keywords",
            "category",
            "comments",
            "last_modified_by",
        ):
            val = getattr(props, attr, None)
            if val:
                meta[attr] = str(val)

        for dt_attr in ("created", "modified"):
            val = getattr(props, dt_attr, None)
            if val:
                meta[dt_attr] = val.isoformat()

        revision = getattr(props, "revision", None)
        if revision is not None:
            meta["revision"] = revision

        return meta
