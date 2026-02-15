"""DOCX report builder for executive reports.

Loads a skeleton DOCX file and programmatically fills it with narrative
text, chart images, and cover page metadata using python-docx. No Jinja2
or docxtpl -- all construction is done via python-docx APIs directly.
"""

import io
import logging
import os
from typing import Any

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Skeleton path resolution
# ---------------------------------------------------------------------------

# Base directory for skeleton files (test-templates/executive/)
_SKELETON_DIR = os.path.join(
    os.path.dirname(__file__),  # sanitization-service/app/services/
    "..", "..", "..",            # -> project root
    "test-templates", "executive",
)

SKELETON_PATHS: dict[str, str] = {
    "en": os.path.normpath(os.path.join(_SKELETON_DIR, "skeleton-en.docx")),
    "pt-pt": os.path.normpath(os.path.join(_SKELETON_DIR, "skeleton-pt-pt.docx")),
    "pt": os.path.normpath(os.path.join(_SKELETON_DIR, "skeleton-pt-pt.docx")),
}


def get_skeleton_path(language: str) -> str:
    """Resolve the skeleton DOCX path for a given language.

    Args:
        language: Language code ("en", "pt-pt", or "pt").

    Returns:
        Absolute path to the skeleton DOCX file.

    Raises:
        FileNotFoundError: If no skeleton exists for the language.
    """
    lang = language.lower()
    path = SKELETON_PATHS.get(lang)
    if not path or not os.path.isfile(path):
        raise FileNotFoundError(
            f"No skeleton DOCX found for language '{language}'. "
            f"Expected at: {path or 'unknown'}"
        )
    return path


# ---------------------------------------------------------------------------
# Section key -> heading text mapping (for flexible matching)
# ---------------------------------------------------------------------------

# EN heading text keyed by section key
_SECTION_HEADINGS_EN: dict[str, str] = {
    "executive_summary": "executive summary",
    "risk_score_explanation": "global risk score",
    "key_metrics_text": "key metrics",
    "severity_analysis": "severity distribution",
    "category_analysis": "vulnerabilities by category",
    "detailed_analysis": "detailed analysis",
    "key_threats": "key threats",
    "compliance_risk_text": "compliance risk",
    "top_vulnerabilities_text": "top 10 vulnerabilities",
    "strategic_recommendations": "strategic recommendations",
    "positive_aspects": "positive aspects",
    "conclusion": "conclusion",
}

# PT heading text keyed by section key
_SECTION_HEADINGS_PT: dict[str, str] = {
    "executive_summary": "sumario executivo",
    "risk_score_explanation": "pontuacao de risco global",
    "key_metrics_text": "metricas principais",
    "severity_analysis": "distribuicao por severidade",
    "category_analysis": "vulnerabilidades por categoria",
    "detailed_analysis": "analise detalhada",
    "key_threats": "principais ameacas",
    "compliance_risk_text": "risco de nao conformidade",
    "top_vulnerabilities_text": "top 10 vulnerabilidades",
    "strategic_recommendations": "recomendacoes estrategicas",
    "positive_aspects": "aspetos positivos",
    "conclusion": "conclusao",
}


class ReportBuilder:
    """Builds executive report DOCX from a skeleton template and data.

    The skeleton DOCX defines branding (headers, footers, fonts, colors,
    logos) and contains section headings and chart placeholder paragraphs
    (e.g., ``[CHART: Severity Distribution]``). This builder fills the
    skeleton with computed narrative text and chart images.
    """

    def __init__(self, skeleton_path: str) -> None:
        """Load the skeleton DOCX into memory.

        Args:
            skeleton_path: Filesystem path to the skeleton DOCX file.
        """
        self._skeleton_path = skeleton_path
        self._doc = Document(skeleton_path)

    def build_report(
        self,
        report_data: dict[str, Any],
        chart_images: dict[str, bytes],
    ) -> bytes:
        """Fill the skeleton DOCX with report content and return bytes.

        Args:
            report_data: Dict with keys:
                - "narrative": dict mapping section_key -> text
                - "metadata": dict with cover page fields
                - "risk_score": float (optional, for score card)
                - "risk_level": str (optional)
            chart_images: Dict mapping chart placeholder names to PNG bytes.
                          Keys are the chart names without brackets, e.g.
                          "Severity Distribution", "Risk Score Card".

        Returns:
            DOCX file as bytes (valid ZIP/DOCX).
        """
        # Work on a fresh copy of the skeleton
        doc = Document(self._skeleton_path)

        # Fill cover page metadata
        metadata = report_data.get("metadata", {})
        self._fill_cover_metadata(doc, metadata)

        # Remove placeholder text paragraphs (the grey "[Section content...]" lines)
        self._remove_placeholder_paragraphs(doc)

        # Fill text sections
        narrative = report_data.get("narrative", {})
        for section_key, text in narrative.items():
            if isinstance(text, str) and text:
                self._fill_text_section(doc, section_key, text)
            elif isinstance(text, dict):
                # Handle nested sections (e.g., strategic_recommendations)
                combined_parts = []
                for sub_key, sub_text in text.items():
                    if isinstance(sub_text, str) and sub_text:
                        combined_parts.append(sub_text)
                if combined_parts:
                    self._fill_text_section(
                        doc, section_key, "\n\n".join(combined_parts)
                    )

        # Replace chart placeholders with images
        for placeholder_name, image_bytes in chart_images.items():
            self._replace_chart_placeholder(doc, placeholder_name, image_bytes)

        # Serialize to bytes
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    def _fill_cover_metadata(self, doc: Document, metadata: dict) -> None:
        """Fill cover page fields (client name, date, project code).

        Searches all paragraphs for placeholder markers like
        ``[CLIENT_NAME]``, ``[PROJECT_CODE]``, ``[REPORT_DATE]``
        and replaces them with actual values.
        """
        replacements = {
            "[CLIENT_NAME]": metadata.get("client_name", ""),
            "[PROJECT_CODE]": metadata.get("project_code", ""),
            "[REPORT_DATE]": metadata.get("report_date", ""),
            "[START_DATE]": metadata.get("start_date", ""),
            "[END_DATE]": metadata.get("end_date", ""),
        }

        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text and value:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)

    def _remove_placeholder_paragraphs(self, doc: Document) -> None:
        """Remove grey placeholder paragraphs from the skeleton.

        These are paragraphs with text like '[Section content will be
        inserted here]' that serve as visual markers in the template
        but should be removed before content is inserted.
        """
        placeholder_text = "[Section content will be inserted here]"
        for paragraph in doc.paragraphs:
            if paragraph.text.strip() == placeholder_text:
                # Remove the paragraph element from the document
                parent = paragraph._element.getparent()
                if parent is not None:
                    parent.remove(paragraph._element)

    def _replace_chart_placeholder(
        self, doc: Document, placeholder_text: str, image_bytes: bytes
    ) -> None:
        """Find paragraph with chart placeholder and replace with image.

        Searches for paragraphs containing ``[CHART: <placeholder_text>]``
        and replaces the paragraph content with the chart image.

        Args:
            doc: The python-docx Document being built.
            placeholder_text: The chart name (e.g., "Severity Distribution").
            image_bytes: PNG image bytes to insert.
        """
        marker = f"[CHART: {placeholder_text}]"

        for paragraph in doc.paragraphs:
            if marker in paragraph.text:
                # Clear existing text
                for run in paragraph.runs:
                    run.text = ""

                # Insert image centered
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                image_stream = io.BytesIO(image_bytes)
                run = paragraph.add_run()
                run.add_picture(image_stream, width=Inches(5.5))
                logger.info("Replaced chart placeholder: %s", marker)
                return

        logger.warning("Chart placeholder not found: %s", marker)

    def _fill_text_section(
        self, doc: Document, section_key: str, text: str
    ) -> None:
        """Find heading matching section key and insert text after it.

        Uses a two-pass matching strategy:
        1. Try exact match against known heading mappings (EN + PT)
        2. Fall back to fuzzy matching (underscores -> spaces, case-insensitive)

        Then removes any existing placeholder paragraphs and inserts the
        narrative text paragraphs after the heading.

        Args:
            doc: The python-docx Document being built.
            section_key: The section identifier (e.g., "executive_summary").
            text: The narrative text to insert.
        """
        target_paragraph = None

        # Pass 1: Match against known heading text for both languages
        # Try all known headings (EN and PT) since the skeleton language
        # may differ from the section key naming convention.
        known_headings = []
        en_heading = _SECTION_HEADINGS_EN.get(section_key)
        pt_heading = _SECTION_HEADINGS_PT.get(section_key)
        if en_heading:
            known_headings.append(en_heading)
        if pt_heading and pt_heading != en_heading:
            known_headings.append(pt_heading)

        for heading in known_headings:
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith("Heading"):
                    if paragraph.text.strip().lower() == heading:
                        target_paragraph = paragraph
                        break
            if target_paragraph:
                break

        # Pass 2: Fuzzy match (original logic)
        if target_paragraph is None:
            search_text = section_key.replace("_", " ").replace(".", " ").lower()
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip().lower()
                if para_text and (search_text in para_text or para_text in search_text):
                    target_paragraph = paragraph
                    break

        if target_paragraph is None:
            logger.warning("Section heading not found for: %s", section_key)
            return

        # Insert text paragraphs after the heading
        # Split by double newlines for paragraph breaks, single newlines for line breaks
        blocks = text.split("\n\n")
        for block in reversed(blocks):
            block = block.strip()
            if not block:
                continue
            # Create paragraph with the text
            new_para = doc.add_paragraph()
            # Handle bold markers (**text**)
            parts = block.split("**")
            for idx, part in enumerate(parts):
                if not part:
                    continue
                run = new_para.add_run(part)
                run.font.size = Pt(11)
                if idx % 2 == 1:  # Odd indices are bold
                    run.bold = True
            # Move paragraph to after the heading
            target_paragraph._element.addnext(new_para._element)

        logger.info("Filled text section: %s", section_key)
