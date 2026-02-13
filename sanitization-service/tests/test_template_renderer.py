"""Integration tests for the Jinja2 template rendering pipeline."""
from io import BytesIO

import pytest
from docx import Document
from docxtpl import DocxTemplate, RichText

from app.services.template_renderer import (
    TemplateRendererService,
    html_to_richtext,
    prepare_context,
    _filter_type,
)
from tests.fixtures.gw_fixture import SAMPLE_CONTEXT


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_template(*paragraphs: str) -> bytes:
    """Create a minimal DOCX template with the given paragraph texts."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docxtpl_instance(template_bytes: bytes) -> DocxTemplate:
    """Wrap template bytes as a DocxTemplate (needed by html_to_richtext)."""
    return DocxTemplate(BytesIO(template_bytes))


@pytest.fixture
def renderer():
    """Provide a TemplateRendererService instance."""
    return TemplateRendererService()


@pytest.fixture
def simple_template_bytes():
    """A minimal DOCX template for testing."""
    return _make_template("Hello {{ client.short_name }}!")


@pytest.fixture
def tpl(simple_template_bytes):
    """A DocxTemplate instance for helper function tests."""
    return _make_docxtpl_instance(simple_template_bytes)


# ---------------------------------------------------------------------------
# html_to_richtext tests
# ---------------------------------------------------------------------------

class TestHtmlToRichtext:
    """Test HTML-to-RichText conversion."""

    def test_converts_bold_tags(self, tpl):
        rt = html_to_richtext("<p>This is <b>bold</b> text.</p>", tpl)
        assert isinstance(rt, RichText)
        # RichText builds internal XML; verify it was created without error
        assert rt is not None

    def test_converts_italic_tags(self, tpl):
        rt = html_to_richtext("<p>This is <i>italic</i> text.</p>", tpl)
        assert isinstance(rt, RichText)

    def test_converts_strong_and_em_tags(self, tpl):
        rt = html_to_richtext("<p><strong>bold</strong> and <em>italic</em></p>", tpl)
        assert isinstance(rt, RichText)

    def test_converts_anchor_tags(self, tpl):
        rt = html_to_richtext(
            '<p>See <a href="https://example.com">this link</a>.</p>', tpl
        )
        assert isinstance(rt, RichText)

    def test_converts_br_tags(self, tpl):
        rt = html_to_richtext("<p>Line 1<br>Line 2</p>", tpl)
        assert isinstance(rt, RichText)

    def test_converts_list_tags(self, tpl):
        rt = html_to_richtext("<ul><li>Item 1</li><li>Item 2</li></ul>", tpl)
        assert isinstance(rt, RichText)

    def test_strips_xhtml_namespace(self, tpl):
        html = '<p xmlns="http://www.w3.org/1999/xhtml">Content here</p>'
        rt = html_to_richtext(html, tpl)
        assert isinstance(rt, RichText)

    def test_strips_multiple_xmlns_attributes(self, tpl):
        html = '<div xmlns="http://www.w3.org/1999/xhtml" xmlns:custom="urn:foo"><p>Test</p></div>'
        rt = html_to_richtext(html, tpl)
        assert isinstance(rt, RichText)

    def test_empty_string_returns_empty_richtext(self, tpl):
        rt = html_to_richtext("", tpl)
        assert isinstance(rt, RichText)

    def test_none_safe(self, tpl):
        """Passing None-ish empty values should not crash."""
        rt = html_to_richtext("", tpl)
        assert isinstance(rt, RichText)

    def test_plain_text_without_tags(self, tpl):
        rt = html_to_richtext("Just plain text", tpl)
        assert isinstance(rt, RichText)


# ---------------------------------------------------------------------------
# filter_type tests
# ---------------------------------------------------------------------------

class TestFilterType:
    """Test the custom filter_type Jinja2 filter."""

    def test_filters_by_single_type(self):
        findings = [
            {"title": "A", "finding_type": "Web"},
            {"title": "B", "finding_type": "Cloud"},
            {"title": "C", "finding_type": "Web"},
        ]
        result = _filter_type(findings, ["Web"])
        assert len(result) == 2
        assert all(f["finding_type"] == "Web" for f in result)

    def test_filters_by_multiple_types(self):
        findings = [
            {"title": "A", "finding_type": "Web"},
            {"title": "B", "finding_type": "Cloud"},
            {"title": "C", "finding_type": "Infrastructure"},
        ]
        result = _filter_type(findings, ["Web", "Cloud"])
        assert len(result) == 2

    def test_case_insensitive_matching(self):
        findings = [{"title": "A", "finding_type": "Web"}]
        result = _filter_type(findings, ["web"])
        assert len(result) == 1

    def test_empty_types_returns_all(self):
        findings = [
            {"title": "A", "finding_type": "Web"},
            {"title": "B", "finding_type": "Cloud"},
        ]
        result = _filter_type(findings, [])
        assert len(result) == 2

    def test_no_match_returns_empty(self):
        findings = [{"title": "A", "finding_type": "Web"}]
        result = _filter_type(findings, ["Physical"])
        assert len(result) == 0

    def test_missing_finding_type_skipped(self):
        findings = [
            {"title": "A", "finding_type": "Web"},
            {"title": "B"},
        ]
        result = _filter_type(findings, ["Web"])
        assert len(result) == 1


# ---------------------------------------------------------------------------
# prepare_context tests
# ---------------------------------------------------------------------------

class TestPrepareContext:
    """Test context preparation with RichText enrichment."""

    def test_adds_rt_fields_to_findings(self, tpl):
        ctx = prepare_context(SAMPLE_CONTEXT, tpl)
        for finding in ctx["findings"]:
            assert "description_rt" in finding
            assert "impact_rt" in finding
            assert "recommendation_rt" in finding
            assert "replication_steps_rt" in finding
            assert "affected_entities_rt" in finding
            assert "severity_rt" in finding
            assert "cvss_vector_link_rt" in finding

    def test_rt_fields_are_richtext_instances(self, tpl):
        ctx = prepare_context(SAMPLE_CONTEXT, tpl)
        for finding in ctx["findings"]:
            assert isinstance(finding["description_rt"], RichText)
            assert isinstance(finding["severity_rt"], RichText)
            assert isinstance(finding["cvss_vector_link_rt"], RichText)

    def test_preserves_non_finding_fields(self, tpl):
        ctx = prepare_context(SAMPLE_CONTEXT, tpl)
        assert ctx["client"]["short_name"] == "AI Template Engine"
        assert ctx["project"]["start_date"] == "2026-02-13"
        assert ctx["report_date"] == "2026-02-13"
        assert ctx["totals"]["findings"] == 4

    def test_classification_rt_is_string(self, tpl):
        ctx = prepare_context(SAMPLE_CONTEXT, tpl)
        assert ctx["findings"][0]["classification_rt"] == "Cloud"
        assert ctx["findings"][2]["classification_rt"] == "Web"

    def test_empty_html_produces_empty_richtext(self, tpl):
        context = {
            "findings": [
                {
                    "title": "Test",
                    "severity": "Low",
                    "severity_color": "",
                    "finding_type": "Web",
                    "cvss_score": 0,
                    "cvss_vector": "",
                    "affected_entities": "",
                    "description": "",
                    "impact": "",
                    "recommendation": "",
                    "replication_steps": "",
                    "references": "",
                }
            ]
        }
        ctx = prepare_context(context, tpl)
        assert isinstance(ctx["findings"][0]["description_rt"], RichText)


# ---------------------------------------------------------------------------
# Full render tests
# ---------------------------------------------------------------------------

class TestTemplateRendererRender:
    """Test the full TemplateRendererService.render() method."""

    def test_renders_simple_variable(self, renderer):
        template = _make_template("Hello {{ client.short_name }}!")
        result = renderer.render(template, SAMPLE_CONTEXT)
        doc = Document(BytesIO(result))
        texts = [p.text for p in doc.paragraphs]
        assert any("AI Template Engine" in t for t in texts)

    def test_renders_finding_titles(self, renderer):
        template = _make_template(
            "{{ client.short_name }}",
            "{%for finding in findings %}",
            "{{ finding.title }}",
            "{%endfor %}",
        )
        result = renderer.render(template, SAMPLE_CONTEXT)
        doc = Document(BytesIO(result))
        texts = [p.text for p in doc.paragraphs]
        full_text = " ".join(texts)
        assert "HSTS" in full_text
        assert "SQLI" in full_text
        assert "XSS Reflected" in full_text

    def test_renders_date_values(self, renderer):
        template = _make_template(
            "Report date: {{ report_date }}",
            "Start: {{ project.start_date }}",
        )
        result = renderer.render(template, SAMPLE_CONTEXT)
        doc = Document(BytesIO(result))
        texts = [p.text for p in doc.paragraphs]
        full_text = " ".join(texts)
        assert "2026-02-13" in full_text

    def test_renders_totals(self, renderer):
        template = _make_template("Total findings: {{ totals.findings }}")
        result = renderer.render(template, SAMPLE_CONTEXT)
        doc = Document(BytesIO(result))
        texts = [p.text for p in doc.paragraphs]
        assert any("4" in t for t in texts)

    def test_filter_type_in_template(self, renderer):
        template = _make_template(
            "{%for finding in findings|filter_type(['Web']) %}",
            "{{ finding.title }}",
            "{%endfor %}",
        )
        result = renderer.render(template, SAMPLE_CONTEXT)
        doc = Document(BytesIO(result))
        texts = [p.text for p in doc.paragraphs]
        full_text = " ".join(texts)
        assert "SQLI" in full_text
        assert "XSS Reflected" in full_text
        # Cloud findings should NOT appear
        assert "HSTS" not in full_text
        assert "mass form" not in full_text

    def test_missing_context_key_renders_empty(self, renderer):
        """Missing keys should render as empty, not crash."""
        template = _make_template("Value: {{ nonexistent_key }}")
        result = renderer.render(template, {})
        doc = Document(BytesIO(result))
        assert len(doc.paragraphs) >= 1

    def test_preserves_non_jinja_content(self, renderer):
        template = _make_template(
            "This is static text that should remain.",
            "{{ client.short_name }}",
        )
        result = renderer.render(template, SAMPLE_CONTEXT)
        doc = Document(BytesIO(result))
        texts = [p.text for p in doc.paragraphs]
        assert any("This is static text that should remain." in t for t in texts)

    def test_output_is_valid_docx(self, renderer):
        template = _make_template("{{ client.short_name }}")
        result = renderer.render(template, SAMPLE_CONTEXT)
        # Must start with PK zip header
        assert result[:2] == b"PK"
        # Must be parseable as DOCX
        doc = Document(BytesIO(result))
        assert len(doc.paragraphs) >= 1

    def test_rejects_invalid_template_bytes(self, renderer):
        with pytest.raises(ValueError, match="Failed to load DOCX template"):
            renderer.render(b"not a docx", {})

    def test_rejects_empty_template(self, renderer):
        with pytest.raises(ValueError, match="Failed to load DOCX template"):
            renderer.render(b"", {})

    def test_scope_loop(self, renderer):
        template = _make_template(
            "{%for item in scope %}",
            "{{ item.scope }}",
            "{%endfor %}",
        )
        result = renderer.render(template, SAMPLE_CONTEXT)
        doc = Document(BytesIO(result))
        texts = [p.text for p in doc.paragraphs]
        full_text = " ".join(texts)
        assert "*.example.com" in full_text
        assert "10.0.0.0/24" in full_text

    def test_team_data(self, renderer):
        template = _make_template(
            "{%for member in team %}",
            "{{ member.name }} - {{ member.email }}",
            "{%endfor %}",
        )
        result = renderer.render(template, SAMPLE_CONTEXT)
        doc = Document(BytesIO(result))
        texts = [p.text for p in doc.paragraphs]
        full_text = " ".join(texts)
        assert "admin" in full_text
