"""Integration tests for adapter API routes."""
import base64
import json
import os
from io import BytesIO

import pytest
from docx import Document as DocxDoc
from fastapi.testclient import TestClient

from app.main import app

TEMPLATES_DIR = os.path.join(
    os.path.dirname(__file__), os.pardir, os.pardir,
    "test-templates", "ghost-templates",
)


@pytest.fixture
def client():
    """FastAPI test client."""
    return TestClient(app)


@pytest.fixture
def web_en_base64():
    """Base64-encoded Web EN reference template."""
    path = os.path.join(TEMPLATES_DIR, "Web_-_EN_2025_-_v2.0_m6w3nHW_FuwLOkd.docx")
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode("ascii")


@pytest.fixture
def minimal_docx_base64():
    """Base64-encoded minimal DOCX for testing."""
    doc = DocxDoc()
    doc.add_heading("Security Assessment Report", level=1)
    doc.add_paragraph("Client Name: Acme Corp")
    doc.add_paragraph("Assessment Period: January 2025 - February 2025")
    doc.add_heading("Findings", level=2)
    doc.add_paragraph("1. SQL Injection vulnerability in login endpoint.")
    buf = BytesIO()
    doc.save(buf)
    return base64.b64encode(buf.getvalue()).decode("ascii")


class TestAnalyzeEndpoint:
    """Tests for POST /adapter/analyze."""

    def test_valid_docx_returns_prompt(self, client, minimal_docx_base64):
        response = client.post(
            "/adapter/analyze",
            json={
                "template_base64": minimal_docx_base64,
                "template_type": "web",
                "language": "en",
            },
        )
        assert response.status_code == 200
        data = response.json()
        assert "prompt" in data
        assert "system_prompt" in data
        assert "reference_template_hash" in data
        assert len(data["prompt"]) > 0
        assert len(data["reference_template_hash"]) == 64  # SHA-256 hex

    def test_returns_doc_structure_summary(self, client, minimal_docx_base64):
        response = client.post(
            "/adapter/analyze",
            json={
                "template_base64": minimal_docx_base64,
                "template_type": "web",
                "language": "en",
            },
        )
        data = response.json()
        assert "doc_structure_summary" in data
        summary = data["doc_structure_summary"]
        assert "paragraph_count" in summary
        assert "non_empty_paragraphs" in summary

    def test_different_template_types(self, client, minimal_docx_base64):
        for ttype in ["web", "internal", "mobile"]:
            response = client.post(
                "/adapter/analyze",
                json={
                    "template_base64": minimal_docx_base64,
                    "template_type": ttype,
                    "language": "en",
                },
            )
            assert response.status_code == 200, f"Failed for type={ttype}"

    def test_rejects_non_docx(self, client):
        # Plain text, not a DOCX
        fake = base64.b64encode(b"This is not a DOCX file").decode("ascii")
        response = client.post(
            "/adapter/analyze",
            json={
                "template_base64": fake,
                "template_type": "web",
                "language": "en",
            },
        )
        assert response.status_code == 400
        assert "not a valid DOCX" in response.json()["detail"]

    def test_rejects_invalid_base64(self, client):
        response = client.post(
            "/adapter/analyze",
            json={
                "template_base64": "!!!not-base64!!!",
                "template_type": "web",
                "language": "en",
            },
        )
        assert response.status_code == 400

    def test_rejects_empty_template(self, client):
        empty = base64.b64encode(b"").decode("ascii")
        response = client.post(
            "/adapter/analyze",
            json={
                "template_base64": empty,
                "template_type": "web",
                "language": "en",
            },
        )
        assert response.status_code == 400

    def test_rejects_invalid_template_type(self, client, minimal_docx_base64):
        response = client.post(
            "/adapter/analyze",
            json={
                "template_base64": minimal_docx_base64,
                "template_type": "desktop",
                "language": "en",
            },
        )
        assert response.status_code == 422  # Pydantic validation error

    # ------ Few-shot integration tests ------

    def test_few_shot_examples_included_in_prompt(self, client, minimal_docx_base64):
        """POST /analyze with few_shot_examples includes the section in prompt."""
        response = client.post(
            "/adapter/analyze",
            json={
                "template_base64": minimal_docx_base64,
                "template_type": "web",
                "language": "en",
                "few_shot_examples": [
                    {
                        "normalized_section_text": "client name: acme corp",
                        "gw_field": "client.short_name",
                        "marker_type": "text",
                        "usage_count": 5,
                    },
                    {
                        "normalized_section_text": "assessment period: january 2025",
                        "gw_field": "project.start_date",
                        "marker_type": "text",
                        "usage_count": 3,
                    },
                    {
                        "normalized_section_text": "detailed vulnerability description",
                        "gw_field": "finding.description_rt",
                        "marker_type": "paragraph_rt",
                        "usage_count": 7,
                    },
                ],
            },
        )
        assert response.status_code == 200
        data = response.json()
        assert "## Previous Successful Mappings" in data["prompt"]
        assert "client name: acme corp" in data["prompt"]
        assert "finding.description_rt" in data["prompt"]

    def test_few_shot_without_field_backward_compat(self, client, minimal_docx_base64):
        """POST /analyze without few_shot_examples field succeeds (backward compat)."""
        response = client.post(
            "/adapter/analyze",
            json={
                "template_base64": minimal_docx_base64,
                "template_type": "web",
                "language": "en",
            },
        )
        assert response.status_code == 200
        data = response.json()
        assert "## Previous Successful Mappings" not in data["prompt"]

    def test_few_shot_empty_array_no_section(self, client, minimal_docx_base64):
        """POST /analyze with empty few_shot_examples produces no few-shot section."""
        response = client.post(
            "/adapter/analyze",
            json={
                "template_base64": minimal_docx_base64,
                "template_type": "web",
                "language": "en",
                "few_shot_examples": [],
            },
        )
        assert response.status_code == 200
        data = response.json()
        assert "## Previous Successful Mappings" not in data["prompt"]

    def test_few_shot_response_structure_intact(self, client, minimal_docx_base64):
        """POST /analyze with few_shot_examples still returns full response structure."""
        response = client.post(
            "/adapter/analyze",
            json={
                "template_base64": minimal_docx_base64,
                "template_type": "web",
                "language": "en",
                "few_shot_examples": [
                    {
                        "normalized_section_text": "test section",
                        "gw_field": "client.short_name",
                        "marker_type": "text",
                        "usage_count": 1,
                    },
                ],
            },
        )
        assert response.status_code == 200
        data = response.json()
        assert "prompt" in data
        assert "system_prompt" in data
        assert "doc_structure_summary" in data
        assert "reference_template_hash" in data
        assert "paragraph_count" in data
        assert len(data["prompt"]) > 0
        assert len(data["reference_template_hash"]) == 64


class TestValidateMappingEndpoint:
    """Tests for POST /adapter/validate-mapping."""

    def test_valid_mapping_returns_plan(self, client):
        llm_response = json.dumps({
            "entries": [
                {
                    "section_index": 0,
                    "section_text": "Client Name: Acme Corp",
                    "gw_field": "client.short_name",
                    "placeholder_template": "{{ client.short_name }}",
                    "confidence": 0.95,
                    "marker_type": "text",
                    "rationale": "Contains client name pattern",
                }
            ],
            "warnings": [],
        })
        response = client.post(
            "/adapter/validate-mapping",
            json={
                "llm_response": llm_response,
                "template_type": "web",
                "language": "en",
                "paragraph_count": 10,
            },
        )
        assert response.status_code == 200
        data = response.json()
        assert data["valid"] is True
        assert data["mapping_plan"] is not None
        assert len(data["mapping_plan"]["entries"]) == 1

    def test_invalid_section_index_returns_error(self, client):
        llm_response = json.dumps({
            "entries": [
                {
                    "section_index": 999,
                    "section_text": "out of range",
                    "gw_field": "client.short_name",
                    "placeholder_template": "{{ client.short_name }}",
                    "confidence": 0.9,
                    "marker_type": "text",
                    "rationale": "test",
                }
            ],
            "warnings": [],
        })
        response = client.post(
            "/adapter/validate-mapping",
            json={
                "llm_response": llm_response,
                "template_type": "web",
                "language": "en",
                "paragraph_count": 10,
            },
        )
        data = response.json()
        assert data["valid"] is False
        assert any("out of range" in e for e in data["errors"])

    def test_invalid_json_returns_error(self, client):
        response = client.post(
            "/adapter/validate-mapping",
            json={
                "llm_response": "this is not json {{{",
                "template_type": "web",
                "language": "en",
                "paragraph_count": 10,
            },
        )
        data = response.json()
        assert data["valid"] is False
        assert any("Invalid JSON" in e for e in data["errors"])

    def test_mismatched_marker_type_returns_error(self, client):
        llm_response = json.dumps({
            "entries": [
                {
                    "section_index": 0,
                    "section_text": "finding desc",
                    "gw_field": "finding.description_rt",
                    "placeholder_template": "{{ finding.description_rt }}",
                    "confidence": 0.9,
                    "marker_type": "text",  # Should be paragraph_rt
                    "rationale": "test",
                }
            ],
            "warnings": [],
        })
        response = client.post(
            "/adapter/validate-mapping",
            json={
                "llm_response": llm_response,
                "template_type": "web",
                "language": "en",
                "paragraph_count": 10,
            },
        )
        data = response.json()
        assert data["valid"] is False
        assert any("marker_type" in e for e in data["errors"])

    def test_multiple_valid_entries(self, client):
        llm_response = json.dumps({
            "entries": [
                {
                    "section_index": 0,
                    "section_text": "Client Name",
                    "gw_field": "client.short_name",
                    "placeholder_template": "{{ client.short_name }}",
                    "confidence": 0.95,
                    "marker_type": "text",
                    "rationale": "name field",
                },
                {
                    "section_index": 2,
                    "section_text": "January 2025",
                    "gw_field": "project.start_date",
                    "placeholder_template": "{{ project.start_date }}",
                    "confidence": 0.85,
                    "marker_type": "text",
                    "rationale": "date field",
                },
            ],
            "warnings": ["Some sections were ambiguous"],
        })
        response = client.post(
            "/adapter/validate-mapping",
            json={
                "llm_response": llm_response,
                "template_type": "web",
                "language": "en",
                "paragraph_count": 10,
            },
        )
        data = response.json()
        assert data["valid"] is True
        assert len(data["mapping_plan"]["entries"]) == 2
        assert len(data["mapping_plan"]["warnings"]) == 1

    def test_negative_section_index_returns_error(self, client):
        llm_response = json.dumps({
            "entries": [
                {
                    "section_index": -1,
                    "section_text": "bad index",
                    "gw_field": "client.short_name",
                    "placeholder_template": "{{ client.short_name }}",
                    "confidence": 0.9,
                    "marker_type": "text",
                    "rationale": "test",
                }
            ],
            "warnings": [],
        })
        response = client.post(
            "/adapter/validate-mapping",
            json={
                "llm_response": llm_response,
                "template_type": "web",
                "language": "en",
                "paragraph_count": 10,
            },
        )
        data = response.json()
        assert data["valid"] is False
        assert any("must be >= 0" in e for e in data["errors"])
