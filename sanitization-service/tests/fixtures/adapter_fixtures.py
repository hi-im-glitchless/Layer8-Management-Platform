"""Reusable test fixtures for template adapter tests.

Provides programmatic DOCX creation and pre-built Pydantic models
(MappingPlan, InstructionSet) for web and internal template types.

All fixtures are realistic (not trivial placeholder data) -- they mirror
actual pentest report structures with corporate styling.
"""
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor

from app.models.adapter import (
    Instruction,
    InstructionSet,
    MappingEntry,
    MappingPlan,
)


# ---------------------------------------------------------------------------
# Programmatic DOCX creation
# ---------------------------------------------------------------------------


def create_test_client_docx() -> bytes:
    """Create a realistic client pentest report DOCX.

    Structure:
    - Cover page: client name + date (paragraphs 0-2)
    - Table of contents placeholder (paragraph 3)
    - Executive summary heading + paragraph (paragraphs 4-5)
    - Scope section with table (paragraphs 6-7 + table)
    - Methodology section (paragraphs 8-9)
    - Findings table (paragraphs 10-11 + table)
    - One sample finding row with text content
    - Footer (paragraph 12)
    - Calibri font throughout, corporate styling

    Returns:
        Raw bytes of the generated DOCX file.
    """
    doc = Document()

    # -- Cover page --
    # Paragraph 0: Title
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("Security Assessment Report")
    r_title.bold = True
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(24)
    r_title.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Paragraph 1: Client name + date
    p_client = doc.add_paragraph()
    r_client = p_client.add_run("Prepared for ACME Corporation")
    r_client.font.name = "Calibri"
    r_client.font.size = Pt(14)

    # Paragraph 2: Report date
    p_date = doc.add_paragraph()
    r_date = p_date.add_run("Report Date: 13 February 2026")
    r_date.font.name = "Calibri"
    r_date.font.size = Pt(11)

    # Paragraph 3: TOC placeholder
    p_toc = doc.add_paragraph()
    r_toc = p_toc.add_run("[Table of Contents]")
    r_toc.font.name = "Calibri"
    r_toc.font.size = Pt(11)
    r_toc.italic = True

    # -- Executive Summary --
    # Paragraph 4: Heading
    doc.add_heading("Executive Summary", level=1)

    # Paragraph 5: Summary text
    p_summary = doc.add_paragraph()
    r_summary = p_summary.add_run(
        "ACME Corporation engaged Layer8 Security to conduct a comprehensive "
        "security assessment of their externally facing web applications. "
        "The assessment was performed between January 2026 and March 2026. "
        "A total of 4 vulnerabilities were identified, ranging from Critical to Low severity."
    )
    r_summary.font.name = "Calibri"
    r_summary.font.size = Pt(11)

    # -- Scope Section --
    # Paragraph 6: Heading
    doc.add_heading("Scope", level=2)

    # Paragraph 7: Scope description
    p_scope_desc = doc.add_paragraph()
    r_scope_desc = p_scope_desc.add_run(
        "The following targets were in scope for this assessment:"
    )
    r_scope_desc.font.name = "Calibri"
    r_scope_desc.font.size = Pt(11)

    # Scope table
    scope_table = doc.add_table(rows=3, cols=2)
    scope_table.rows[0].cells[0].text = "Target"
    scope_table.rows[0].cells[1].text = "Type"
    scope_table.rows[1].cells[0].text = "*.acme-corp.com"
    scope_table.rows[1].cells[1].text = "Web Application"
    scope_table.rows[2].cells[0].text = "10.0.0.0/24"
    scope_table.rows[2].cells[1].text = "Internal Network"

    # -- Methodology --
    # Paragraph 8: Heading
    doc.add_heading("Methodology", level=2)

    # Paragraph 9: Methodology text
    p_method = doc.add_paragraph()
    r_method = p_method.add_run(
        "The assessment followed the OWASP Testing Guide v4.2 methodology "
        "and included both automated scanning and manual testing techniques."
    )
    r_method.font.name = "Calibri"
    r_method.font.size = Pt(11)

    # -- Findings --
    # Paragraph 10: Heading
    doc.add_heading("Findings", level=1)

    # Paragraph 11: Findings intro
    p_findings_intro = doc.add_paragraph()
    r_findings_intro = p_findings_intro.add_run(
        "The following table summarises the vulnerabilities identified during the assessment."
    )
    r_findings_intro.font.name = "Calibri"
    r_findings_intro.font.size = Pt(11)

    # Findings table with headers + one sample row
    findings_table = doc.add_table(rows=2, cols=5)
    headers = ["Title", "Severity", "Description", "Impact", "Recommendation"]
    for i, header in enumerate(headers):
        cell = findings_table.rows[0].cells[i]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True

    findings_table.rows[1].cells[0].text = "SQL Injection in Login"
    findings_table.rows[1].cells[1].text = "Critical"
    findings_table.rows[1].cells[2].text = "Unsanitized user input in the login form allows SQL injection attacks."
    findings_table.rows[1].cells[3].text = "Full database compromise including user credentials."
    findings_table.rows[1].cells[4].text = "Use parameterized queries for all database interactions."

    # Paragraph 12: Footer/team info
    p_team = doc.add_paragraph()
    r_team = p_team.add_run("Lead Assessor: Jane Smith (jane@layer8.local)")
    r_team.font.name = "Calibri"
    r_team.font.size = Pt(11)

    # Add page numbers via a footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    r_footer = footer_para.add_run("Page ")
    r_footer.font.name = "Calibri"
    r_footer.font.size = Pt(9)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Pre-built MappingPlan fixtures
# ---------------------------------------------------------------------------


SAMPLE_MAPPING_PLAN_WEB = MappingPlan(
    entries=[
        MappingEntry(
            section_index=1,
            section_text="Prepared for ACME Corporation",
            gw_field="client.short_name",
            placeholder_template="{{ client.short_name }}",
            confidence=0.95,
            marker_type="text",
            rationale="Client name in cover page",
        ),
        MappingEntry(
            section_index=2,
            section_text="Report Date: 13 February 2026",
            gw_field="report_date",
            placeholder_template="{{ report_date }}",
            confidence=0.90,
            marker_type="text",
            rationale="Report date on cover page",
        ),
        MappingEntry(
            section_index=5,
            section_text="between January 2026 and March 2026",
            gw_field="project.start_date",
            placeholder_template="{{ project.start_date }}",
            confidence=0.85,
            marker_type="text",
            rationale="Project start date in executive summary",
        ),
        MappingEntry(
            section_index=5,
            section_text="total of 4 vulnerabilities",
            gw_field="totals.findings",
            placeholder_template="{{ totals.findings }}",
            confidence=0.80,
            marker_type="text",
            rationale="Finding count in executive summary",
        ),
        MappingEntry(
            section_index=12,
            section_text="Jane Smith",
            gw_field="team[0].name",
            placeholder_template="{{ team[0].name }}",
            confidence=0.90,
            marker_type="text",
            rationale="Lead assessor name in team section",
        ),
    ],
    template_type="web",
    language="en",
    warnings=["Scope table may need manual review for loop placement"],
)


SAMPLE_MAPPING_PLAN_INTERNAL = MappingPlan(
    entries=[
        MappingEntry(
            section_index=1,
            section_text="Prepared for ACME Corporation",
            gw_field="client.short_name",
            placeholder_template="{{ client.short_name }}",
            confidence=0.95,
            marker_type="text",
            rationale="Client name in cover page",
        ),
        MappingEntry(
            section_index=2,
            section_text="Report Date: 13 February 2026",
            gw_field="report_date",
            placeholder_template="{{ report_date }}",
            confidence=0.90,
            marker_type="text",
            rationale="Report date on cover page",
        ),
        MappingEntry(
            section_index=5,
            section_text="between January 2026 and March 2026",
            gw_field="project.start_date",
            placeholder_template="{{ project.start_date }}",
            confidence=0.85,
            marker_type="text",
            rationale="Project date range",
        ),
    ],
    template_type="internal",
    language="en",
    warnings=["Internal template may need filter_type adjustments"],
)


# ---------------------------------------------------------------------------
# Pre-built InstructionSet fixtures
# ---------------------------------------------------------------------------


SAMPLE_INSTRUCTION_SET_WEB = InstructionSet(
    instructions=[
        Instruction(
            action="replace_text",
            paragraph_index=1,
            original_text="ACME Corporation",
            replacement_text="{{ client.short_name }}",
            marker_type="text",
            gw_field="client.short_name",
        ),
        Instruction(
            action="replace_text",
            paragraph_index=2,
            original_text="13 February 2026",
            replacement_text="{{ report_date }}",
            marker_type="text",
            gw_field="report_date",
        ),
        Instruction(
            action="replace_text",
            paragraph_index=5,
            original_text="January 2026",
            replacement_text="{{ project.start_date }}",
            marker_type="text",
            gw_field="project.start_date",
        ),
        Instruction(
            action="replace_text",
            paragraph_index=12,
            original_text="Jane Smith",
            replacement_text="{{ team[0].name }}",
            marker_type="text",
            gw_field="team[0].name",
        ),
    ],
    template_type="web",
    language="en",
)


SAMPLE_INSTRUCTION_SET_INTERNAL = InstructionSet(
    instructions=[
        Instruction(
            action="replace_text",
            paragraph_index=1,
            original_text="ACME Corporation",
            replacement_text="{{ client.short_name }}",
            marker_type="text",
            gw_field="client.short_name",
        ),
        Instruction(
            action="replace_text",
            paragraph_index=2,
            original_text="13 February 2026",
            replacement_text="{{ report_date }}",
            marker_type="text",
            gw_field="report_date",
        ),
        Instruction(
            action="replace_text",
            paragraph_index=5,
            original_text="January 2026",
            replacement_text="{{ project.start_date }}",
            marker_type="text",
            gw_field="project.start_date",
        ),
    ],
    template_type="internal",
    language="en",
    additional_blocks=[
        '{% for finding in findings|filter_type(["Web"]) %}',
        '{% endfor %}',
        '{% set ns = namespace(counter=0) %}',
    ],
)
