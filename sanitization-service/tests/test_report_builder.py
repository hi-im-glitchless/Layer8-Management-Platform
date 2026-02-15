"""Unit tests for report_builder.py -- DOCX report builder."""

import io
import os
import tempfile
import zipfile

import pytest
from docx import Document

from app.services.report_builder import ReportBuilder


# ZIP/DOCX signature bytes (PK header)
DOCX_SIGNATURE = b"PK"


@pytest.fixture
def skeleton_path():
    """Create a minimal skeleton DOCX for testing and return its path."""
    doc = Document()

    # Cover page placeholders
    doc.add_paragraph("[CLIENT_NAME]")
    doc.add_paragraph("[PROJECT_CODE]")
    doc.add_paragraph("[REPORT_DATE]")

    # Section headings
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph("Placeholder executive summary text.")

    doc.add_heading("Risk Score Explanation", level=1)
    doc.add_paragraph("Placeholder risk text.")

    doc.add_heading("Severity Analysis", level=1)
    doc.add_paragraph("[CHART: Severity Distribution]")

    doc.add_heading("Category Analysis", level=1)
    doc.add_paragraph("[CHART: Category Distribution]")

    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph("Placeholder conclusion.")

    # Write to temp file
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc.save(f)
        path = f.name

    yield path

    # Cleanup
    os.unlink(path)


@pytest.fixture
def sample_chart_png():
    """Create a minimal valid PNG for chart placeholder testing."""
    # Minimal 1x1 transparent PNG
    import struct
    import zlib

    def _make_png():
        signature = b"\x89PNG\r\n\x1a\n"

        # IHDR
        width = 1
        height = 1
        bit_depth = 8
        color_type = 2  # RGB
        ihdr_data = struct.pack(">IIBBBBB", width, height, bit_depth, color_type, 0, 0, 0)
        ihdr_crc = zlib.crc32(b"IHDR" + ihdr_data) & 0xFFFFFFFF
        ihdr = struct.pack(">I", 13) + b"IHDR" + ihdr_data + struct.pack(">I", ihdr_crc)

        # IDAT
        raw_data = b"\x00\xff\x00\x00"  # filter byte + RGB
        compressed = zlib.compress(raw_data)
        idat_crc = zlib.crc32(b"IDAT" + compressed) & 0xFFFFFFFF
        idat = struct.pack(">I", len(compressed)) + b"IDAT" + compressed + struct.pack(">I", idat_crc)

        # IEND
        iend_crc = zlib.crc32(b"IEND") & 0xFFFFFFFF
        iend = struct.pack(">I", 0) + b"IEND" + struct.pack(">I", iend_crc)

        return signature + ihdr + idat + iend

    return _make_png()


# ---------------------------------------------------------------------------
# ReportBuilder init tests
# ---------------------------------------------------------------------------


class TestReportBuilderInit:
    def test_loads_skeleton(self, skeleton_path):
        builder = ReportBuilder(skeleton_path)
        assert builder._doc is not None

    def test_invalid_path_raises(self):
        with pytest.raises(Exception):
            ReportBuilder("/nonexistent/path.docx")


# ---------------------------------------------------------------------------
# build_report tests
# ---------------------------------------------------------------------------


class TestBuildReport:
    def test_returns_valid_docx_bytes(self, skeleton_path):
        builder = ReportBuilder(skeleton_path)
        result = builder.build_report(
            report_data={"metadata": {}, "narrative": {}},
            chart_images={},
        )
        assert isinstance(result, bytes)
        assert result[:2] == DOCX_SIGNATURE

    def test_docx_is_valid_zip(self, skeleton_path):
        builder = ReportBuilder(skeleton_path)
        result = builder.build_report(
            report_data={"metadata": {}, "narrative": {}},
            chart_images={},
        )
        buf = io.BytesIO(result)
        assert zipfile.is_zipfile(buf)

    def test_docx_contains_content_types(self, skeleton_path):
        """Valid DOCX must contain [Content_Types].xml."""
        builder = ReportBuilder(skeleton_path)
        result = builder.build_report(
            report_data={"metadata": {}, "narrative": {}},
            chart_images={},
        )
        buf = io.BytesIO(result)
        with zipfile.ZipFile(buf) as zf:
            assert "[Content_Types].xml" in zf.namelist()

    def test_fills_metadata_placeholders(self, skeleton_path):
        builder = ReportBuilder(skeleton_path)
        result = builder.build_report(
            report_data={
                "metadata": {
                    "client_name": "Acme Corp",
                    "project_code": "PT-2025-042",
                    "report_date": "2025-02-15",
                },
                "narrative": {},
            },
            chart_images={},
        )
        # Parse the output DOCX and check for replaced values
        doc = Document(io.BytesIO(result))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Acme Corp" in all_text
        assert "PT-2025-042" in all_text
        assert "2025-02-15" in all_text

    def test_fills_narrative_sections(self, skeleton_path):
        builder = ReportBuilder(skeleton_path)
        result = builder.build_report(
            report_data={
                "metadata": {},
                "narrative": {
                    "executive_summary": "The assessment revealed critical issues.",
                    "conclusion": "In summary, immediate action is required.",
                },
            },
            chart_images={},
        )
        doc = Document(io.BytesIO(result))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "critical issues" in all_text
        assert "immediate action is required" in all_text

    def test_replaces_chart_placeholders(self, skeleton_path, sample_chart_png):
        builder = ReportBuilder(skeleton_path)
        result = builder.build_report(
            report_data={"metadata": {}, "narrative": {}},
            chart_images={"Severity Distribution": sample_chart_png},
        )
        doc = Document(io.BytesIO(result))
        # The placeholder text should be gone
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "[CHART: Severity Distribution]" not in all_text

    def test_empty_report_data(self, skeleton_path):
        """Empty data should still produce valid DOCX."""
        builder = ReportBuilder(skeleton_path)
        result = builder.build_report(
            report_data={},
            chart_images={},
        )
        assert result[:2] == DOCX_SIGNATURE

    def test_handles_nested_recommendations(self, skeleton_path):
        builder = ReportBuilder(skeleton_path)
        result = builder.build_report(
            report_data={
                "metadata": {},
                "narrative": {
                    "strategic_recommendations": {
                        "immediate": "Patch critical vulnerabilities.",
                        "short_term": "Deploy WAF.",
                    },
                },
            },
            chart_images={},
        )
        assert result[:2] == DOCX_SIGNATURE
