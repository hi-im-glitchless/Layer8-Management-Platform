"""End-to-end template rendering smoke test.

Loads the actual Web EN reference template, renders it with static GW fixture
data, and verifies the output is a valid DOCX containing expected content.

This test proves the full pipeline works: DOCX template + GW data -> rendered DOCX.
It does NOT require a running Ghostwriter instance (uses static fixture data)
or Gotenberg (tests DOCX output only, not PDF conversion).
"""
import os
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from app.services.template_renderer import TemplateRendererService
from tests.fixtures.gw_fixture import SAMPLE_CONTEXT

# Path to the reference template relative to the project root
_PROJECT_ROOT = Path(__file__).resolve().parents[2]
_TEMPLATE_PATH = (
    _PROJECT_ROOT
    / "test-templates"
    / "ghost-templates"
    / "Web_-_EN_2025_-_v2.0_m6w3nHW_FuwLOkd.docx"
)

# Mark as integration test -- depends on the reference template file existing
pytestmark = pytest.mark.integration


def _template_exists() -> bool:
    return _TEMPLATE_PATH.is_file()


@pytest.fixture(scope="module")
def renderer():
    return TemplateRendererService()


@pytest.fixture(scope="module")
def template_bytes():
    if not _template_exists():
        pytest.skip(f"Reference template not found: {_TEMPLATE_PATH}")
    return _TEMPLATE_PATH.read_bytes()


@pytest.fixture(scope="module")
def rendered_bytes(renderer, template_bytes):
    """Render the reference template once for all tests in this module."""
    return renderer.render(template_bytes, SAMPLE_CONTEXT)


@pytest.fixture(scope="module")
def rendered_doc(rendered_bytes):
    """Parse the rendered DOCX for content inspection."""
    return Document(BytesIO(rendered_bytes))


@pytest.fixture(scope="module")
def rendered_text(rendered_doc) -> str:
    """Extract all text from the rendered document (paragraphs + tables)."""
    texts: list[str] = []

    for para in rendered_doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)

    for table in rendered_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    texts.append(cell.text)

    return "\n".join(texts)


class TestE2ERenderOutputValidity:
    """Verify the rendered output is a valid, non-trivial DOCX."""

    def test_output_is_valid_docx(self, rendered_bytes):
        """Rendered output must be a valid DOCX (parseable by python-docx)."""
        assert rendered_bytes[:2] == b"PK", "Output should start with PK zip header"
        doc = Document(BytesIO(rendered_bytes))
        assert len(doc.paragraphs) > 0

    def test_output_has_content(self, rendered_doc):
        """Rendered document should have substantial content."""
        assert len(rendered_doc.paragraphs) > 10, "Document should have many paragraphs"

    def test_output_has_tables(self, rendered_doc):
        """Reference template uses tables for findings; rendered doc should too."""
        assert len(rendered_doc.tables) > 0, "Document should contain tables"


class TestE2ERenderClientData:
    """Verify client and project data appears in the rendered output."""

    def test_client_name_appears(self, rendered_text):
        assert "AI Template Engine" in rendered_text, (
            "Client short_name should appear in rendered document"
        )

    def test_report_date_appears(self, rendered_text):
        assert "2026-02-13" in rendered_text, (
            "Report date should appear in rendered document"
        )

    def test_project_start_date_appears(self, rendered_text):
        assert "2026-02-13" in rendered_text, (
            "Project start date should appear in rendered document"
        )

    def test_project_end_date_appears(self, rendered_text):
        assert "2026-12-11" in rendered_text, (
            "Project end date should appear in rendered document"
        )


class TestE2ERenderFindings:
    """Verify finding data appears in the rendered output."""

    def test_finding_title_hsts(self, rendered_text):
        assert "HSTS" in rendered_text, "Finding title 'HSTS' should appear"

    def test_finding_title_sqli(self, rendered_text):
        assert "SQLI" in rendered_text, "Finding title 'SQLI' should appear"

    def test_finding_title_xss(self, rendered_text):
        assert "XSS Reflected" in rendered_text, "Finding title 'XSS Reflected' should appear"

    def test_finding_title_mass_form(self, rendered_text):
        assert "mass form" in rendered_text, "Finding title 'mass form' should appear"

    def test_finding_count(self, rendered_text):
        """Total findings count (4) should appear somewhere."""
        assert "4" in rendered_text, "Total findings count should appear"


class TestE2ERenderTeam:
    """Verify team data appears in the rendered output."""

    def test_team_member_name(self, rendered_text):
        assert "admin" in rendered_text, "Team member name 'admin' should appear"
