"""Unit tests for annotated preview service (paragraph shading + metadata)."""
import base64
from io import BytesIO
from unittest.mock import patch

import pytest
from docx import Document as DocxDoc
from docx.oxml.ns import qn
from fastapi.testclient import TestClient

from app.main import app
from app.models.adapter import MappingEntry, MappingPlan
from app.models.gap_detection import GapEntry
from app.services.annotated_preview import (
    GREEN_SHADING,
    YELLOW_SHADING,
    apply_paragraph_shading,
    generate_annotation_metadata,
)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


def _make_test_docx(paragraphs: list[str] | None = None) -> bytes:
    """Create a minimal test DOCX with the given paragraphs."""
    doc = DocxDoc()
    if paragraphs is None:
        paragraphs = [
            "Security Assessment Report",
            "Client Name: Acme Corp",
            "Assessment Period: January 2025",
            "",  # empty paragraph
            "Finding: SQL Injection",
        ]
    for text in paragraphs:
        if text.startswith("# "):
            doc.add_heading(text[2:], level=1)
        elif text.startswith("## "):
            doc.add_heading(text[3:], level=2)
        else:
            doc.add_paragraph(text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_mapping_entry(
    gw_field: str,
    section_index: int,
    marker_type: str = "text",
    section_text: str = "",
) -> MappingEntry:
    return MappingEntry(
        section_index=section_index,
        section_text=section_text or f"Text at index {section_index}",
        gw_field=gw_field,
        placeholder_template=f"{{{{ {gw_field} }}}}",
        confidence=0.9,
        marker_type=marker_type,
    )


@pytest.fixture
def test_docx_bytes():
    """5-paragraph test DOCX."""
    return _make_test_docx()


@pytest.fixture
def mapping_plan():
    """Mapping plan with 2 entries for the test DOCX."""
    return MappingPlan(
        entries=[
            _make_mapping_entry("client.short_name", 1, section_text="Client Name: Acme Corp"),
            _make_mapping_entry("finding.title", 4, section_text="Finding: SQL Injection"),
        ],
        template_type="web",
        language="en",
    )


@pytest.fixture
def gap_entries():
    """Gap entries for paragraph index 2."""
    return [
        GapEntry(
            gw_field="project.start_date",
            marker_type="text",
            expected_context="Assessment period reference",
            estimated_paragraph_index=2,
        ),
    ]


@pytest.fixture
def client():
    """FastAPI test client."""
    return TestClient(app)


# ---------------------------------------------------------------------------
# Paragraph shading tests
# ---------------------------------------------------------------------------


class TestApplyParagraphShading:
    """Tests for apply_paragraph_shading()."""

    def test_green_shading_on_mapped_paragraphs(self, test_docx_bytes, mapping_plan):
        """Mapped paragraphs should have green background shading."""
        result_bytes = apply_paragraph_shading(test_docx_bytes, mapping_plan, [])
        doc = DocxDoc(BytesIO(result_bytes))

        # Paragraph at index 1 (Client Name) should have green shading
        para = doc.paragraphs[1]
        shd = para._element.find(f".//{qn('w:shd')}")
        assert shd is not None, "Expected w:shd element on mapped paragraph"
        assert shd.get(qn("w:fill")) == GREEN_SHADING

    def test_yellow_shading_on_gap_paragraphs(self, test_docx_bytes, mapping_plan, gap_entries):
        """Gap paragraphs should have yellow background shading."""
        result_bytes = apply_paragraph_shading(test_docx_bytes, mapping_plan, gap_entries)
        doc = DocxDoc(BytesIO(result_bytes))

        # Paragraph at index 2 (Assessment Period) should have yellow shading
        para = doc.paragraphs[2]
        shd = para._element.find(f".//{qn('w:shd')}")
        assert shd is not None, "Expected w:shd element on gap paragraph"
        assert shd.get(qn("w:fill")) == YELLOW_SHADING

    def test_shading_xml_attributes(self, test_docx_bytes, mapping_plan):
        """Shading XML should have val=clear, color=auto, fill=hex."""
        result_bytes = apply_paragraph_shading(test_docx_bytes, mapping_plan, [])
        doc = DocxDoc(BytesIO(result_bytes))

        para = doc.paragraphs[1]
        shd = para._element.find(f".//{qn('w:shd')}")
        assert shd is not None
        assert shd.get(qn("w:val")) == "clear"
        assert shd.get(qn("w:color")) == "auto"
        assert shd.get(qn("w:fill")) == GREEN_SHADING

    def test_unshaded_paragraphs_unchanged(self, test_docx_bytes, mapping_plan):
        """Paragraphs not in the mapping plan or gaps should have no shading."""
        result_bytes = apply_paragraph_shading(test_docx_bytes, mapping_plan, [])
        doc = DocxDoc(BytesIO(result_bytes))

        # Paragraph at index 0 (heading) should have no shading added
        para = doc.paragraphs[0]
        shd = para._element.find(f".//{qn('w:shd')}")
        # May or may not have shd from existing style; check no C6EFCE or FFF2CC
        if shd is not None:
            fill = shd.get(qn("w:fill"))
            assert fill not in (GREEN_SHADING, YELLOW_SHADING)

    def test_mapped_takes_priority_over_gap(self, test_docx_bytes):
        """If a paragraph is both mapped and a gap candidate, green wins."""
        plan = MappingPlan(
            entries=[_make_mapping_entry("client.short_name", 1)],
            template_type="web",
            language="en",
        )
        gaps = [
            GapEntry(
                gw_field="project.start_date",
                marker_type="text",
                expected_context="",
                estimated_paragraph_index=1,  # Same index as mapped
            ),
        ]

        result_bytes = apply_paragraph_shading(test_docx_bytes, plan, gaps)
        doc = DocxDoc(BytesIO(result_bytes))

        para = doc.paragraphs[1]
        shd = para._element.find(f".//{qn('w:shd')}")
        assert shd is not None
        assert shd.get(qn("w:fill")) == GREEN_SHADING  # Green, not yellow

    def test_out_of_range_index_skipped(self, test_docx_bytes):
        """Out-of-range section_index should be skipped without error."""
        plan = MappingPlan(
            entries=[_make_mapping_entry("client.short_name", 99)],
            template_type="web",
            language="en",
        )
        # Should not raise
        result_bytes = apply_paragraph_shading(test_docx_bytes, plan, [])
        assert len(result_bytes) > 0


# ---------------------------------------------------------------------------
# Annotation metadata tests
# ---------------------------------------------------------------------------


class TestGenerateAnnotationMetadata:
    """Tests for generate_annotation_metadata()."""

    def test_tooltip_entries_for_mapped_paragraphs(self, test_docx_bytes, mapping_plan):
        """Mapped entries should appear as tooltip entries with status='mapped'."""
        metadata = generate_annotation_metadata(test_docx_bytes, mapping_plan, [])

        mapped_tooltips = [t for t in metadata.tooltip_data if t.status == "mapped"]
        assert len(mapped_tooltips) == 2

        fields = {t.gw_field for t in mapped_tooltips}
        assert "client.short_name" in fields
        assert "finding.title" in fields

    def test_tooltip_entries_for_gap_paragraphs(self, test_docx_bytes, mapping_plan, gap_entries):
        """Gap entries should appear as tooltip entries with status='gap'."""
        metadata = generate_annotation_metadata(test_docx_bytes, mapping_plan, gap_entries)

        gap_tooltips = [t for t in metadata.tooltip_data if t.status == "gap"]
        assert len(gap_tooltips) == 1
        assert gap_tooltips[0].gw_field == "project.start_date"
        assert gap_tooltips[0].paragraph_index == 2

    def test_unmapped_paragraphs_excludes_empty(self, test_docx_bytes, mapping_plan):
        """Empty paragraphs should not appear in unmapped_paragraphs."""
        metadata = generate_annotation_metadata(test_docx_bytes, mapping_plan, [])

        # Paragraph index 3 is empty -- should not be in unmapped list
        unmapped_indices = {u.paragraph_index for u in metadata.unmapped_paragraphs}
        assert 3 not in unmapped_indices

    def test_unmapped_paragraphs_excludes_mapped(self, test_docx_bytes, mapping_plan):
        """Mapped paragraphs should not appear in unmapped_paragraphs."""
        metadata = generate_annotation_metadata(test_docx_bytes, mapping_plan, [])

        unmapped_indices = {u.paragraph_index for u in metadata.unmapped_paragraphs}
        assert 1 not in unmapped_indices  # client.short_name mapped
        assert 4 not in unmapped_indices  # finding.title mapped

    def test_unmapped_paragraphs_sorted_by_index(self, test_docx_bytes, mapping_plan):
        """Unmapped paragraphs should be sorted by paragraph_index ascending."""
        metadata = generate_annotation_metadata(test_docx_bytes, mapping_plan, [])

        indices = [u.paragraph_index for u in metadata.unmapped_paragraphs]
        assert indices == sorted(indices)

    def test_unmapped_text_truncated(self):
        """Unmapped paragraph text should be truncated to 200 chars."""
        long_text = "A" * 500
        doc_bytes = _make_test_docx([long_text])

        plan = MappingPlan(
            entries=[],
            template_type="web",
            language="en",
        )

        metadata = generate_annotation_metadata(doc_bytes, plan, [])

        assert len(metadata.unmapped_paragraphs) == 1
        assert len(metadata.unmapped_paragraphs[0].text) == 200


# ---------------------------------------------------------------------------
# Annotate endpoint integration tests
# ---------------------------------------------------------------------------


class TestAnnotateEndpoint:
    """Tests for POST /adapter/annotate endpoint."""

    @patch("app.routes.adapter.detect_gaps")
    def test_annotate_returns_annotated_base64(self, mock_detect, client):
        """POST /annotate should return annotated_base64 and metadata."""
        from app.models.gap_detection import GapDetectionResult

        mock_detect.return_value = GapDetectionResult(
            gaps=[],
            mapped_field_count=1,
            expected_field_count=1,
            coverage_percent=100.0,
        )

        doc_bytes = _make_test_docx()
        b64 = base64.b64encode(doc_bytes).decode("ascii")

        response = client.post(
            "/adapter/annotate",
            json={
                "template_base64": b64,
                "mapping_plan": {
                    "entries": [
                        {
                            "section_index": 1,
                            "section_text": "Client Name: Acme Corp",
                            "gw_field": "client.short_name",
                            "placeholder_template": "{{ client.short_name }}",
                            "confidence": 0.9,
                            "marker_type": "text",
                        },
                    ],
                    "template_type": "web",
                    "language": "en",
                },
            },
        )

        assert response.status_code == 200
        data = response.json()
        assert "annotated_base64" in data
        assert "tooltip_data" in data
        assert "unmapped_paragraphs" in data
        assert "gap_summary" in data

        # Verify annotated_base64 is valid base64 DOCX
        annotated_bytes = base64.b64decode(data["annotated_base64"])
        assert annotated_bytes[:4] == b"PK\x03\x04"

    def test_annotate_invalid_base64_returns_400(self, client):
        """Invalid base64 should return 400."""
        response = client.post(
            "/adapter/annotate",
            json={
                "template_base64": "not-valid-base64!!!",
                "mapping_plan": {
                    "entries": [],
                    "template_type": "web",
                    "language": "en",
                },
            },
        )
        assert response.status_code == 400

    def test_annotate_non_docx_returns_400(self, client):
        """Non-DOCX content should return 400."""
        b64 = base64.b64encode(b"this is not a docx file").decode("ascii")
        response = client.post(
            "/adapter/annotate",
            json={
                "template_base64": b64,
                "mapping_plan": {
                    "entries": [],
                    "template_type": "web",
                    "language": "en",
                },
            },
        )
        assert response.status_code == 400

    @patch("app.routes.adapter.detect_gaps")
    def test_annotate_includes_gap_summary(self, mock_detect, client):
        """Response should include gap_summary with coverage data."""
        from app.models.gap_detection import GapDetectionResult

        mock_detect.return_value = GapDetectionResult(
            gaps=[],
            mapped_field_count=3,
            expected_field_count=5,
            coverage_percent=60.0,
        )

        doc_bytes = _make_test_docx()
        b64 = base64.b64encode(doc_bytes).decode("ascii")

        response = client.post(
            "/adapter/annotate",
            json={
                "template_base64": b64,
                "mapping_plan": {
                    "entries": [],
                    "template_type": "web",
                    "language": "en",
                },
            },
        )

        assert response.status_code == 200
        gap_summary = response.json()["gap_summary"]
        assert gap_summary["mapped_field_count"] == 3
        assert gap_summary["expected_field_count"] == 5
        assert gap_summary["coverage_percent"] == 60.0

    @patch("app.routes.adapter.detect_gaps")
    def test_annotate_endpoint_green_only_param(self, mock_detect, client):
        """POST /annotate with green_only=true should return annotated DOCX without yellow shading."""
        from app.models.gap_detection import GapDetectionResult, GapEntry as GapEntryModel

        mock_detect.return_value = GapDetectionResult(
            gaps=[
                GapEntryModel(
                    gw_field="project.start_date",
                    marker_type="text",
                    expected_context="Assessment period reference",
                    estimated_paragraph_index=2,
                ),
            ],
            mapped_field_count=1,
            expected_field_count=2,
            coverage_percent=50.0,
        )

        doc_bytes = _make_test_docx()
        b64 = base64.b64encode(doc_bytes).decode("ascii")

        response = client.post(
            "/adapter/annotate",
            json={
                "template_base64": b64,
                "mapping_plan": {
                    "entries": [
                        {
                            "section_index": 1,
                            "section_text": "Client Name: Acme Corp",
                            "gw_field": "client.short_name",
                            "placeholder_template": "{{ client.short_name }}",
                            "confidence": 0.9,
                            "marker_type": "text",
                        },
                    ],
                    "template_type": "web",
                    "language": "en",
                },
                "green_only": True,
            },
        )

        assert response.status_code == 200
        data = response.json()

        # Verify the annotated DOCX is returned
        annotated_bytes = base64.b64decode(data["annotated_base64"])
        doc = DocxDoc(BytesIO(annotated_bytes))

        # Mapped paragraph (index 1) should have green shading
        para1 = doc.paragraphs[1]
        shd1 = para1._element.find(f".//{qn('w:shd')}")
        assert shd1 is not None
        assert shd1.get(qn("w:fill")) == GREEN_SHADING

        # Gap paragraph (index 2) should NOT have yellow shading
        para2 = doc.paragraphs[2]
        shd2 = para2._element.find(f".//{qn('w:shd')}")
        if shd2 is not None:
            fill = shd2.get(qn("w:fill"))
            assert fill != YELLOW_SHADING, "Gap paragraph should not have yellow shading in green_only mode"


# ---------------------------------------------------------------------------
# Green-only shading tests
# ---------------------------------------------------------------------------


class TestGreenOnlyShading:
    """Tests for green_only mode in apply_paragraph_shading()."""

    def test_apply_paragraph_shading_green_only(self, test_docx_bytes, mapping_plan, gap_entries):
        """With green_only=True, only mapped paragraphs get green shading, gaps have none."""
        result_bytes = apply_paragraph_shading(
            test_docx_bytes, mapping_plan, gap_entries, green_only=True
        )
        doc = DocxDoc(BytesIO(result_bytes))

        # Paragraph at index 1 (Client Name) should have green shading
        para1 = doc.paragraphs[1]
        shd1 = para1._element.find(f".//{qn('w:shd')}")
        assert shd1 is not None, "Expected green shading on mapped paragraph"
        assert shd1.get(qn("w:fill")) == GREEN_SHADING

        # Paragraph at index 4 (Finding: SQL Injection) should also have green shading
        para4 = doc.paragraphs[4]
        shd4 = para4._element.find(f".//{qn('w:shd')}")
        assert shd4 is not None, "Expected green shading on second mapped paragraph"
        assert shd4.get(qn("w:fill")) == GREEN_SHADING

        # Paragraph at index 2 (gap candidate) should NOT have any shading
        para2 = doc.paragraphs[2]
        shd2 = para2._element.find(f".//{qn('w:shd')}")
        if shd2 is not None:
            fill = shd2.get(qn("w:fill"))
            assert fill not in (YELLOW_SHADING, GREEN_SHADING), (
                f"Gap paragraph should not have shading in green_only mode, got fill={fill}"
            )

    def test_apply_paragraph_shading_default_includes_yellow(
        self, test_docx_bytes, mapping_plan, gap_entries
    ):
        """Default behavior (green_only not specified) should include yellow gap shading."""
        result_bytes = apply_paragraph_shading(test_docx_bytes, mapping_plan, gap_entries)
        doc = DocxDoc(BytesIO(result_bytes))

        # Paragraph at index 2 (gap) should have yellow shading
        para2 = doc.paragraphs[2]
        shd2 = para2._element.find(f".//{qn('w:shd')}")
        assert shd2 is not None, "Expected yellow shading on gap paragraph by default"
        assert shd2.get(qn("w:fill")) == YELLOW_SHADING

        # Paragraph at index 1 (mapped) should still have green
        para1 = doc.paragraphs[1]
        shd1 = para1._element.find(f".//{qn('w:shd')}")
        assert shd1 is not None
        assert shd1.get(qn("w:fill")) == GREEN_SHADING
