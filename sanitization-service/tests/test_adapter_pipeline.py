"""Integration tests for the full template adapter pipeline.

Tests the complete flow: build instruction set -> enrich via rules_engine ->
validate via jinja2_validator -> apply via InstructionApplier -> verify output DOCX.

All test DOCX files are created programmatically -- no external services needed.
"""
from io import BytesIO

import pytest
from docx import Document
from docx.shared import Pt, RGBColor

from app.models.adapter import (
    Instruction,
    InstructionSet,
    MappingEntry,
    MappingPlan,
)
from app.services.instruction_applier import InstructionApplier
from app.services.jinja2_validator import validate_instruction_set
from app.services.rules_engine import enrich_instructions


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _create_client_docx() -> bytes:
    """Create a realistic client DOCX programmatically.

    Contains:
    - Heading "Executive Summary" (bold)
    - Paragraph "Client Name: ACME Corp" (Calibri 11pt)
    - Table with headers ["Finding", "Severity", "Description"] + one data row
    - Paragraph "Project dates: January 2026 - March 2026"
    - Paragraph "Testing scope: www.example.com"
    """
    doc = Document()

    # Heading
    heading = doc.add_heading("Executive Summary", level=1)

    # Client name paragraph
    p_client = doc.add_paragraph()
    run = p_client.add_run("Client Name: ACME Corp")
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    # Findings table
    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "Finding"
    table.rows[0].cells[1].text = "Severity"
    table.rows[0].cells[2].text = "Description"
    table.rows[1].cells[0].text = "SQL Injection"
    table.rows[1].cells[1].text = "Critical"
    table.rows[1].cells[2].text = "Unsanitized user input in SQL query"

    # Project dates
    p_dates = doc.add_paragraph()
    run_dates = p_dates.add_run("Project dates: January 2026 - March 2026")
    run_dates.font.name = "Calibri"
    run_dates.font.size = Pt(11)

    # Scope
    p_scope = doc.add_paragraph()
    run_scope = p_scope.add_run("Testing scope: www.example.com")
    run_scope.font.name = "Calibri"
    run_scope.font.size = Pt(11)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _parse_docx(docx_bytes: bytes) -> Document:
    """Parse DOCX bytes back into a Document."""
    return Document(BytesIO(docx_bytes))


def _make_instruction(
    action: str = "replace_text",
    paragraph_index: int = 0,
    original_text: str = "",
    replacement_text: str = "",
    marker_type: str = "text",
    gw_field: str = "",
) -> Instruction:
    return Instruction(
        action=action,
        paragraph_index=paragraph_index,
        original_text=original_text,
        replacement_text=replacement_text,
        marker_type=marker_type,
        gw_field=gw_field,
    )


def _make_instruction_set(
    instructions: list[Instruction],
    template_type: str = "web",
    language: str = "en",
    additional_blocks: list[str] | None = None,
) -> InstructionSet:
    return InstructionSet(
        instructions=instructions,
        template_type=template_type,
        language=language,
        additional_blocks=additional_blocks or [],
    )


# ---------------------------------------------------------------------------
# Pipeline Tests: Web Template
# ---------------------------------------------------------------------------


class TestPipelineWebTemplate:
    """Integration test: full pipeline for web template type."""

    def test_pipeline_web_template(self):
        """Build instructions, enrich, validate, apply for web template."""
        docx_bytes = _create_client_docx()

        # Build instruction set from a synthetic mapping plan.
        # Each paragraph_index + action pair must be unique (validator enforces this).
        # Paragraph layout: [0]=heading, [1]=client name, [2]=dates, [3]=scope
        instructions = [
            _make_instruction(
                action="replace_text",
                paragraph_index=1,  # "Client Name: ACME Corp"
                original_text="ACME Corp",
                replacement_text="{{ client.short_name }}",
                marker_type="text",
                gw_field="client.short_name",
            ),
            _make_instruction(
                action="replace_text",
                paragraph_index=2,  # "Project dates: January 2026 - March 2026"
                original_text="January 2026 - March 2026",
                replacement_text="{{ project.start_date }} - {{ project.end_date }}",
                marker_type="text",
                gw_field="project.start_date",
            ),
            _make_instruction(
                action="replace_text",
                paragraph_index=3,  # "Testing scope: www.example.com"
                original_text="www.example.com",
                replacement_text="{{ item.scope }}",
                marker_type="text",
                gw_field="item.scope",
            ),
        ]
        iset = _make_instruction_set(instructions, template_type="web", language="en")

        # Enrich via rules engine
        enriched = enrich_instructions(iset)

        # Validate via jinja2_validator
        validation = validate_instruction_set(enriched)
        assert validation.valid, f"Validation failed: {validation.errors}"

        # Apply via InstructionApplier
        applier = InstructionApplier()
        safe_iset = validation.sanitized_instructions or enriched
        output_bytes, applied, skipped, warnings = applier.apply(docx_bytes, safe_iset)

        # Verify output DOCX
        out_doc = _parse_docx(output_bytes)
        assert applied == 3, f"Expected 3 applied, got {applied}"

        # Check client.short_name was inserted
        p1_text = out_doc.paragraphs[1].text
        assert "{{ client.short_name }}" in p1_text
        assert "ACME Corp" not in p1_text

        # Check dates were replaced
        p2_text = out_doc.paragraphs[2].text
        assert "{{ project.start_date }}" in p2_text
        assert "{{ project.end_date }}" in p2_text

        # Check scope was replaced
        p3_text = out_doc.paragraphs[3].text
        assert "{{ item.scope }}" in p3_text

        # Verify formatting preserved on paragraph 1
        runs = out_doc.paragraphs[1].runs
        for run in runs:
            if run.font.name is not None:
                assert run.font.name == "Calibri"
            if run.font.size is not None:
                assert run.font.size == Pt(11)

        # Heading style should be preserved
        assert out_doc.paragraphs[0].style.name.startswith("Heading")

        # Enrichment should have added additional blocks for web template
        assert len(enriched.additional_blocks) > 0


# ---------------------------------------------------------------------------
# Pipeline Tests: Internal Template
# ---------------------------------------------------------------------------


class TestPipelineInternalTemplate:
    """Integration test: full pipeline for internal template type."""

    def test_pipeline_internal_template(self):
        """Internal template adds filter_type and namespace counters."""
        docx_bytes = _create_client_docx()

        instructions = [
            _make_instruction(
                action="replace_text",
                paragraph_index=1,
                original_text="ACME Corp",
                replacement_text="{{ client.short_name }}",
                marker_type="text",
                gw_field="client.short_name",
            ),
            _make_instruction(
                action="replace_text",
                paragraph_index=2,
                original_text="January 2026",
                replacement_text="{{ project.start_date }}",
                marker_type="text",
                gw_field="project.start_date",
            ),
        ]
        iset = _make_instruction_set(instructions, template_type="internal", language="en")

        # Enrich via rules engine
        enriched = enrich_instructions(iset)

        # Verify filter_type blocks injected for internal
        additional_text = " ".join(enriched.additional_blocks)
        assert "filter_type" in additional_text, "Internal templates must have filter_type blocks"

        # Verify namespace counters injected
        assert "namespace(counter=0)" in additional_text, "Internal templates must have namespace counters"

        # Validate
        validation = validate_instruction_set(enriched)
        assert validation.valid, f"Validation failed: {validation.errors}"

        # Apply
        applier = InstructionApplier()
        safe_iset = validation.sanitized_instructions or enriched
        output_bytes, applied, skipped, warnings = applier.apply(docx_bytes, safe_iset)

        assert applied >= 2
        out_doc = _parse_docx(output_bytes)
        assert "{{ client.short_name }}" in out_doc.paragraphs[1].text


# ---------------------------------------------------------------------------
# Pipeline Tests: Mobile Template
# ---------------------------------------------------------------------------


class TestPipelineMobileTemplate:
    """Integration test: full pipeline for mobile template type."""

    def test_pipeline_mobile_template(self):
        """Mobile template uses scope loops and affected_entities."""
        docx_bytes = _create_client_docx()

        instructions = [
            _make_instruction(
                action="replace_text",
                paragraph_index=1,
                original_text="ACME Corp",
                replacement_text="{{ client.short_name }}",
                marker_type="text",
                gw_field="client.short_name",
            ),
            _make_instruction(
                action="replace_text",
                paragraph_index=3,  # scope paragraph
                original_text="www.example.com",
                replacement_text="{{ item.scope }}",
                marker_type="text",
                gw_field="item.scope",
            ),
        ]
        iset = _make_instruction_set(instructions, template_type="mobile", language="en")

        # Enrich via rules engine
        enriched = enrich_instructions(iset)

        # Verify scope loops and affected_entities injected for mobile
        additional_text = " ".join(enriched.additional_blocks)
        assert "scope" in additional_text, "Mobile templates must have scope loop blocks"
        assert "affected_entities" in additional_text, "Mobile templates must have affected_entities blocks"

        # Validate
        validation = validate_instruction_set(enriched)
        assert validation.valid, f"Validation failed: {validation.errors}"

        # Apply
        applier = InstructionApplier()
        safe_iset = validation.sanitized_instructions or enriched
        output_bytes, applied, skipped, warnings = applier.apply(docx_bytes, safe_iset)

        assert applied >= 2
        out_doc = _parse_docx(output_bytes)
        assert "{{ client.short_name }}" in out_doc.paragraphs[1].text
        assert "{{ item.scope }}" in out_doc.paragraphs[3].text


# ---------------------------------------------------------------------------
# Pipeline Tests: Invalid Instructions Rejected
# ---------------------------------------------------------------------------


class TestPipelineInvalidInstructionsRejected:
    """Verify that unsafe Jinja2 is rejected before reaching the applier."""

    def test_pipeline_invalid_instructions_rejected(self):
        """Instructions with unsafe Jinja2 are rejected by the validator."""
        docx_bytes = _create_client_docx()

        # Build instructions with dangerous Jinja2
        instructions = [
            _make_instruction(
                action="replace_text",
                paragraph_index=1,
                original_text="ACME Corp",
                replacement_text="{{ ''.__class__.__mro__ }}",
                marker_type="text",
                gw_field="client.short_name",
            ),
        ]
        iset = _make_instruction_set(instructions, template_type="web", language="en")

        # Enrich first (should not filter dangerous patterns)
        enriched = enrich_instructions(iset)

        # Validate should reject
        validation = validate_instruction_set(enriched)
        assert not validation.valid, "Validator must reject unsafe Jinja2"
        assert len(validation.errors) > 0
        assert any("Dangerous" in e or "__" in e for e in validation.errors)

        # The sanitized instructions should be empty (no valid instructions)
        safe_iset = validation.sanitized_instructions
        assert safe_iset is not None
        assert len(safe_iset.instructions) == 0, "No instructions should survive validation"

    def test_pipeline_import_os_rejected(self):
        """Instructions with import statements are rejected."""
        instructions = [
            _make_instruction(
                action="replace_text",
                paragraph_index=0,
                original_text="test",
                replacement_text="{{ import('os').system('rm -rf /') }}",
                marker_type="text",
                gw_field="client.short_name",
            ),
        ]
        iset = _make_instruction_set(instructions)
        validation = validate_instruction_set(iset)
        assert not validation.valid
        assert any("Dangerous" in e or "import" in e for e in validation.errors)

    def test_pipeline_mixed_valid_and_invalid(self):
        """Mix of valid and invalid instructions: only valid survive."""
        instructions = [
            # Valid
            _make_instruction(
                action="replace_text",
                paragraph_index=0,
                original_text="test",
                replacement_text="{{ client.short_name }}",
                marker_type="text",
                gw_field="client.short_name",
            ),
            # Invalid
            _make_instruction(
                action="replace_text",
                paragraph_index=1,
                original_text="test2",
                replacement_text="{{ ''.__class__.__mro__ }}",
                marker_type="text",
                gw_field="client.short_name",
            ),
        ]
        iset = _make_instruction_set(instructions)
        validation = validate_instruction_set(iset)
        assert not validation.valid  # Overall invalid because of errors
        assert validation.sanitized_instructions is not None
        assert len(validation.sanitized_instructions.instructions) == 1  # Only the valid one


# ---------------------------------------------------------------------------
# Pipeline Tests: Formatting Preservation
# ---------------------------------------------------------------------------


class TestPipelinePreservesFormatting:
    """Verify formatting is preserved through the full pipeline."""

    def test_pipeline_preserves_formatting(self):
        """Create DOCX with specific formatting, apply instructions, verify preserved."""
        doc = Document()

        # Red bold text
        p1 = doc.add_paragraph()
        r1 = p1.add_run("Critical Finding Here")
        r1.bold = True
        r1.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        r1.font.name = "Calibri"
        r1.font.size = Pt(12)

        # 14pt heading style
        heading = doc.add_heading("Findings Section", level=2)

        # Blue italic text
        p3 = doc.add_paragraph()
        r3 = p3.add_run("Assessment Period: Q1 2026")
        r3.italic = True
        r3.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
        r3.font.name = "Arial"
        r3.font.size = Pt(10)

        buf = BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        # Build instructions that replace text within formatted runs
        instructions = [
            _make_instruction(
                action="replace_text",
                paragraph_index=0,
                original_text="Critical Finding Here",
                replacement_text="{{ finding.title }}",
                marker_type="text",
                gw_field="finding.title",
            ),
            _make_instruction(
                action="replace_text",
                paragraph_index=2,
                original_text="Q1 2026",
                replacement_text="{{ project.start_date }}",
                marker_type="text",
                gw_field="project.start_date",
            ),
        ]
        iset = _make_instruction_set(instructions, template_type="web", language="en")

        # Full pipeline: enrich -> validate -> apply
        enriched = enrich_instructions(iset)
        validation = validate_instruction_set(enriched)
        assert validation.valid, f"Validation failed: {validation.errors}"

        applier = InstructionApplier()
        safe_iset = validation.sanitized_instructions or enriched
        output_bytes, applied, skipped, warnings = applier.apply(docx_bytes, safe_iset)
        assert applied == 2

        # Verify formatting preserved
        out_doc = _parse_docx(output_bytes)

        # Paragraph 0: red bold Calibri 12pt
        p0_runs = out_doc.paragraphs[0].runs
        assert len(p0_runs) >= 1
        r0 = p0_runs[0]
        assert "{{ finding.title }}" in r0.text
        assert r0.bold is True
        assert str(r0.font.color.rgb) == "FF0000"
        assert r0.font.name == "Calibri"
        assert r0.font.size == Pt(12)

        # Paragraph 1: heading style preserved
        assert out_doc.paragraphs[1].style.name.startswith("Heading")

        # Paragraph 2: blue italic Arial 10pt
        p2_runs = out_doc.paragraphs[2].runs
        # Find the run containing our replacement
        found_replacement = False
        for run in p2_runs:
            if "{{ project.start_date }}" in run.text:
                found_replacement = True
                # The italic flag should be preserved on the run that contained the original text
                assert run.italic is True
                assert str(run.font.color.rgb) == "0000FF"
                assert run.font.name == "Arial"
                assert run.font.size == Pt(10)
                break
        assert found_replacement, "Expected to find replacement text in paragraph 2 runs"

    def test_preserves_underline(self):
        """Verify underline formatting survives the pipeline."""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("Underlined Client Name")
        r.underline = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)

        buf = BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        instructions = [
            _make_instruction(
                action="replace_text",
                paragraph_index=0,
                original_text="Underlined Client Name",
                replacement_text="{{ client.short_name }}",
                marker_type="text",
                gw_field="client.short_name",
            ),
        ]
        iset = _make_instruction_set(instructions)
        enriched = enrich_instructions(iset)
        validation = validate_instruction_set(enriched)
        assert validation.valid

        applier = InstructionApplier()
        output_bytes, applied, _, _ = applier.apply(docx_bytes, validation.sanitized_instructions or enriched)
        assert applied == 1

        out_doc = _parse_docx(output_bytes)
        out_run = out_doc.paragraphs[0].runs[0]
        assert out_run.underline is True
        assert out_run.font.name == "Times New Roman"
        assert out_run.font.size == Pt(14)
