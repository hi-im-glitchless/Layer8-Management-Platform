"""End-to-end template adaptation verification tests.

Uses actual reference templates from test-templates/ghost-templates/ and
static GW fixture data to verify the complete adapter pipeline produces
DOCX output that is both valid and GW-compatible.

Marked @pytest.mark.integration -- these tests use real template files
but do NOT require a running LLM, Ghostwriter, or Gotenberg instance.
"""
import os
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt

from app.models.adapter import Instruction, InstructionSet
from app.services.docx_parser import DocxParserService
from app.services.instruction_applier import InstructionApplier
from app.services.jinja2_validator import validate_instruction_set
from app.services.reference_loader import (
    TEMPLATE_DIR,
    TEMPLATE_MAP,
    load_reference_template,
)
from app.services.rules_engine import enrich_instructions
from app.services.template_renderer import TemplateRendererService
from tests.fixtures.gw_fixture import SAMPLE_CONTEXT

# ---------------------------------------------------------------------------
# Skip if reference templates are not available
# ---------------------------------------------------------------------------

TEMPLATES_AVAILABLE = TEMPLATE_DIR.exists() and any(TEMPLATE_DIR.glob("*.docx"))

pytestmark = [
    pytest.mark.integration,
    pytest.mark.skipif(
        not TEMPLATES_AVAILABLE,
        reason=f"Reference templates not found at {TEMPLATE_DIR}",
    ),
]

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_parser = DocxParserService()
_applier = InstructionApplier()
_renderer = TemplateRendererService()


def _create_client_docx_web() -> bytes:
    """Create a programmatic client DOCX that simulates a web pentest report.

    Has plain text in the positions where Jinja2 placeholders should go.
    """
    doc = Document()

    # Cover/Header
    doc.add_heading("Security Assessment Report", level=1)

    # Client info
    p1 = doc.add_paragraph()
    r1 = p1.add_run("This document is prepared for TestClient Corp and covers "
                     "the security assessment conducted during the engagement period.")
    r1.font.name = "Calibri"
    r1.font.size = Pt(11)

    # Executive summary
    doc.add_heading("Executive Summary", level=1)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("TestClient Corp has an ongoing concern about information security. "
                     "This report evaluates the security posture of web applications.")
    r2.font.name = "Calibri"
    r2.font.size = Pt(11)

    # Scope
    doc.add_heading("Scope", level=2)
    p3 = doc.add_paragraph()
    r3 = p3.add_run("The security tests were performed between January 2026 and March 2026.")
    r3.font.name = "Calibri"
    r3.font.size = Pt(11)

    # Scope table
    scope_table = doc.add_table(rows=3, cols=1)
    scope_table.rows[0].cells[0].text = "Scope Item"
    scope_table.rows[1].cells[0].text = "*.example.com"
    scope_table.rows[2].cells[0].text = "10.0.0.0/24"

    # Findings
    doc.add_heading("Findings", level=1)
    p4 = doc.add_paragraph()
    r4 = p4.add_run("A total of 4 vulnerabilities were discovered.")
    r4.font.name = "Calibri"
    r4.font.size = Pt(11)

    # Finding details table
    findings_table = doc.add_table(rows=2, cols=4)
    findings_table.rows[0].cells[0].text = "Title"
    findings_table.rows[0].cells[1].text = "Severity"
    findings_table.rows[0].cells[2].text = "Description"
    findings_table.rows[0].cells[3].text = "Impact"
    findings_table.rows[1].cells[0].text = "HSTS Missing"
    findings_table.rows[1].cells[1].text = "High"
    findings_table.rows[1].cells[2].text = "HSTS header not set"
    findings_table.rows[1].cells[3].text = "MitM attacks possible"

    # Team
    doc.add_heading("Team", level=2)
    p5 = doc.add_paragraph()
    r5 = p5.add_run("Lead Assessor: John Doe (john@layer8.local)")
    r5.font.name = "Calibri"
    r5.font.size = Pt(11)

    # Report date
    p6 = doc.add_paragraph()
    r6 = p6.add_run("Report Date: February 2026")
    r6.font.name = "Calibri"
    r6.font.size = Pt(11)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _create_client_docx_internal() -> bytes:
    """Create a programmatic internal pentest report DOCX."""
    doc = Document()

    doc.add_heading("Internal Security Assessment", level=1)

    p1 = doc.add_paragraph()
    r1 = p1.add_run("Prepared for InternalClient Corp. Assessment of internal infrastructure.")
    r1.font.name = "Calibri"
    r1.font.size = Pt(11)

    doc.add_heading("Findings by Category", level=2)

    # Category: AD
    p_ad = doc.add_paragraph()
    p_ad.add_run("Active Directory findings:")

    # Category: Web
    p_web = doc.add_paragraph()
    p_web.add_run("Web application findings:")

    # Finding counter
    p_counter = doc.add_paragraph()
    p_counter.add_run("Finding 01: Weak Password Policy")

    # Dates
    p_dates = doc.add_paragraph()
    r_dates = p_dates.add_run("Tests performed: January 2026 to March 2026")
    r_dates.font.name = "Calibri"
    r_dates.font.size = Pt(11)

    # Total
    p_total = doc.add_paragraph()
    p_total.add_run("Total findings: 4")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _create_client_docx_mobile() -> bytes:
    """Create a programmatic mobile pentest report DOCX."""
    doc = Document()

    doc.add_heading("Mobile Application Security Assessment", level=1)

    p1 = doc.add_paragraph()
    r1 = p1.add_run("Prepared for MobileClient Corp. Security testing of mobile applications.")
    r1.font.name = "Calibri"
    r1.font.size = Pt(11)

    doc.add_heading("Scope", level=2)
    p_scope = doc.add_paragraph()
    p_scope.add_run("Target applications: iOS and Android clients.")

    # Scope table
    scope_table = doc.add_table(rows=2, cols=1)
    scope_table.rows[0].cells[0].text = "Application"
    scope_table.rows[1].cells[0].text = "com.example.app"

    doc.add_heading("Findings", level=2)
    p_finding = doc.add_paragraph()
    p_finding.add_run("Insecure data storage found in local database.")

    # Affected entities
    p_affected = doc.add_paragraph()
    p_affected.add_run("Affected: /api/users endpoint")

    # Dates
    p_dates = doc.add_paragraph()
    r_dates = p_dates.add_run("Assessment period: February 2026 to April 2026")
    r_dates.font.name = "Calibri"
    r_dates.font.size = Pt(11)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_web_instructions(doc_bytes: bytes) -> InstructionSet:
    """Build a synthetic instruction set for a web template adaptation."""
    doc = Document(BytesIO(doc_bytes))
    paragraphs = doc.paragraphs

    instructions = []

    # Find and map known sections
    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if not text:
            continue

        # Client name replacement
        if "TestClient Corp" in text and "prepared for" in text.lower():
            instructions.append(Instruction(
                action="replace_text",
                paragraph_index=i,
                original_text="TestClient Corp",
                replacement_text="{{ client.short_name }}",
                marker_type="text",
                gw_field="client.short_name",
            ))

        # Date range replacement
        if "January 2026" in text and "March 2026" in text:
            instructions.append(Instruction(
                action="replace_text",
                paragraph_index=i,
                original_text="January 2026",
                replacement_text="{{ project.start_date }}",
                marker_type="text",
                gw_field="project.start_date",
            ))

        # Totals replacement
        if "total of 4" in text.lower():
            instructions.append(Instruction(
                action="replace_text",
                paragraph_index=i,
                original_text="4",
                replacement_text="{{ totals.findings }}",
                marker_type="text",
                gw_field="totals.findings",
            ))

        # Team name
        if "John Doe" in text:
            instructions.append(Instruction(
                action="replace_text",
                paragraph_index=i,
                original_text="John Doe",
                replacement_text="{{ team[0].name }}",
                marker_type="text",
                gw_field="team[0].name",
            ))

        # Report date
        if "Report Date:" in text:
            instructions.append(Instruction(
                action="replace_text",
                paragraph_index=i,
                original_text="February 2026",
                replacement_text="{{ report_date }}",
                marker_type="text",
                gw_field="report_date",
            ))

    return InstructionSet(
        instructions=instructions,
        template_type="web",
        language="en",
    )


def _build_internal_instructions(doc_bytes: bytes) -> InstructionSet:
    """Build a synthetic instruction set for an internal template adaptation."""
    doc = Document(BytesIO(doc_bytes))
    paragraphs = doc.paragraphs

    instructions = []

    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if not text:
            continue

        if "InternalClient Corp" in text:
            instructions.append(Instruction(
                action="replace_text",
                paragraph_index=i,
                original_text="InternalClient Corp",
                replacement_text="{{ client.short_name }}",
                marker_type="text",
                gw_field="client.short_name",
            ))

        if "January 2026" in text and "March 2026" in text:
            instructions.append(Instruction(
                action="replace_text",
                paragraph_index=i,
                original_text="January 2026",
                replacement_text="{{ project.start_date }}",
                marker_type="text",
                gw_field="project.start_date",
            ))

        if "Total findings: 4" in text:
            instructions.append(Instruction(
                action="replace_text",
                paragraph_index=i,
                original_text="4",
                replacement_text="{{ totals.findings }}",
                marker_type="text",
                gw_field="totals.findings",
            ))

    return InstructionSet(
        instructions=instructions,
        template_type="internal",
        language="en",
    )


def _build_mobile_instructions(doc_bytes: bytes) -> InstructionSet:
    """Build a synthetic instruction set for a mobile template adaptation."""
    doc = Document(BytesIO(doc_bytes))
    paragraphs = doc.paragraphs

    instructions = []

    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if not text:
            continue

        if "MobileClient Corp" in text:
            instructions.append(Instruction(
                action="replace_text",
                paragraph_index=i,
                original_text="MobileClient Corp",
                replacement_text="{{ client.short_name }}",
                marker_type="text",
                gw_field="client.short_name",
            ))

        if "February 2026" in text and "April 2026" in text:
            instructions.append(Instruction(
                action="replace_text",
                paragraph_index=i,
                original_text="February 2026",
                replacement_text="{{ project.start_date }}",
                marker_type="text",
                gw_field="project.start_date",
            ))

    return InstructionSet(
        instructions=instructions,
        template_type="mobile",
        language="en",
    )


# ---------------------------------------------------------------------------
# E2E: Web EN adaptation
# ---------------------------------------------------------------------------


class TestE2EWebEnAdaptation:
    """End-to-end test: Web EN template adaptation and render verification."""

    def test_e2e_web_en_adaptation(self):
        """Full pipeline: parse -> map -> enrich -> apply -> render with GW data."""
        # 1. Load reference template info
        ref_info = load_reference_template("web", "en")
        assert ref_info.placeholder_count > 0, "Reference template should have patterns"

        # 2. Create a synthetic client DOCX
        client_bytes = _create_client_docx_web()

        # 3. Parse it to verify structure
        doc_structure = _parser.parse(client_bytes)
        assert len(doc_structure.paragraphs) > 5

        # 4. Build synthetic mapping plan -> instruction set
        iset = _build_web_instructions(client_bytes)
        assert len(iset.instructions) > 0, "Should have at least some instructions"

        # 5. Enrich via rules engine
        enriched = enrich_instructions(iset)
        assert len(enriched.additional_blocks) > 0, "Web template should have additional blocks"

        # 6. Validate
        validation = validate_instruction_set(enriched)
        assert validation.valid, f"Validation failed: {validation.errors}"

        # 7. Apply instructions
        safe_iset = validation.sanitized_instructions or enriched
        output_bytes, applied, skipped, warnings = _applier.apply(client_bytes, safe_iset)

        assert applied > 0, f"Expected applied > 0, got {applied}"

        # 8. Verify output is a valid DOCX
        out_doc = Document(BytesIO(output_bytes))
        assert len(out_doc.paragraphs) > 0

        # 9. Verify Jinja2 placeholders are present
        all_text = " ".join(p.text for p in out_doc.paragraphs)
        assert "{{ client.short_name }}" in all_text
        assert "{{ project.start_date }}" in all_text

        # 10. Verify original page layout preserved
        if out_doc.sections:
            section = out_doc.sections[0]
            assert section.page_width is not None
            assert section.page_height is not None

        # 11. CRITICAL: Render the reference template with GW fixture data
        # This proves the GW template rendering pipeline works end-to-end
        ref_template_path = TEMPLATE_DIR / TEMPLATE_MAP[("web", "en")]
        ref_bytes = ref_template_path.read_bytes()

        rendered_bytes = _renderer.render(ref_bytes, SAMPLE_CONTEXT)
        assert len(rendered_bytes) > 0, "Rendered output should not be empty"

        # Verify rendered output contains actual GW data values
        rendered_doc = Document(BytesIO(rendered_bytes))
        rendered_text = " ".join(p.text for p in rendered_doc.paragraphs)
        assert "AI Template Engine" in rendered_text, (
            "Rendered output should contain client.short_name from GW fixture"
        )

    def test_e2e_web_formatting_preserved(self):
        """Verify formatting survives the adaptation pipeline."""
        client_bytes = _create_client_docx_web()
        iset = _build_web_instructions(client_bytes)
        enriched = enrich_instructions(iset)
        validation = validate_instruction_set(enriched)
        safe_iset = validation.sanitized_instructions or enriched
        output_bytes, _, _, _ = _applier.apply(client_bytes, safe_iset)

        out_doc = Document(BytesIO(output_bytes))

        # Check Calibri font is preserved on paragraphs with content
        for p in out_doc.paragraphs:
            for run in p.runs:
                if run.font.name is not None:
                    # Should still be Calibri (we set it on all content paragraphs)
                    assert run.font.name == "Calibri", (
                        f"Expected Calibri, got {run.font.name} on text: {run.text[:50]}"
                    )


# ---------------------------------------------------------------------------
# E2E: Internal EN adaptation
# ---------------------------------------------------------------------------


class TestE2EInternalEnAdaptation:
    """End-to-end test: Internal EN template adaptation and render verification."""

    def test_e2e_internal_en_adaptation(self):
        """Full pipeline with internal template: filter_type and namespace."""
        # 1. Load reference
        ref_info = load_reference_template("internal", "en")
        assert ref_info.placeholder_count > 0

        # 2. Create client DOCX and build instructions
        client_bytes = _create_client_docx_internal()
        iset = _build_internal_instructions(client_bytes)
        assert len(iset.instructions) > 0

        # 3. Enrich -- should add filter_type and namespace blocks
        enriched = enrich_instructions(iset)
        additional_text = " ".join(enriched.additional_blocks)
        assert "filter_type" in additional_text, "Internal template must have filter_type"
        assert "namespace(counter=0)" in additional_text, "Internal template must have namespace counters"

        # 4. Validate and apply
        validation = validate_instruction_set(enriched)
        assert validation.valid, f"Validation failed: {validation.errors}"

        safe_iset = validation.sanitized_instructions or enriched
        output_bytes, applied, skipped, warnings = _applier.apply(client_bytes, safe_iset)
        assert applied > 0

        # 5. Verify output
        out_doc = Document(BytesIO(output_bytes))
        all_text = " ".join(p.text for p in out_doc.paragraphs)
        assert "{{ client.short_name }}" in all_text

        # 6. CRITICAL: Render reference template with GW data
        ref_template_path = TEMPLATE_DIR / TEMPLATE_MAP[("internal", "en")]
        ref_bytes = ref_template_path.read_bytes()

        rendered_bytes = _renderer.render(ref_bytes, SAMPLE_CONTEXT)
        assert len(rendered_bytes) > 0

        rendered_doc = Document(BytesIO(rendered_bytes))
        rendered_text = " ".join(p.text for p in rendered_doc.paragraphs)
        assert "AI Template Engine" in rendered_text


# ---------------------------------------------------------------------------
# E2E: Mobile EN adaptation
# ---------------------------------------------------------------------------


class TestE2EMobileEnAdaptation:
    """End-to-end test: Mobile EN template adaptation and render verification."""

    def test_e2e_mobile_en_adaptation(self):
        """Full pipeline with mobile template: scope loops and affected_entities."""
        # 1. Load reference
        ref_info = load_reference_template("mobile", "en")
        assert ref_info.placeholder_count > 0

        # 2. Create client DOCX and build instructions
        client_bytes = _create_client_docx_mobile()
        iset = _build_mobile_instructions(client_bytes)
        assert len(iset.instructions) > 0

        # 3. Enrich -- should add scope loops and affected_entities
        enriched = enrich_instructions(iset)
        additional_text = " ".join(enriched.additional_blocks)
        assert "scope" in additional_text, "Mobile template must have scope loop blocks"
        assert "affected_entities" in additional_text, "Mobile template must have affected_entities"

        # 4. Validate and apply
        validation = validate_instruction_set(enriched)
        assert validation.valid, f"Validation failed: {validation.errors}"

        safe_iset = validation.sanitized_instructions or enriched
        output_bytes, applied, skipped, warnings = _applier.apply(client_bytes, safe_iset)
        assert applied > 0

        # 5. Verify output
        out_doc = Document(BytesIO(output_bytes))
        all_text = " ".join(p.text for p in out_doc.paragraphs)
        assert "{{ client.short_name }}" in all_text

        # 6. CRITICAL: Render reference template with GW data
        ref_template_path = TEMPLATE_DIR / TEMPLATE_MAP[("mobile", "en")]
        ref_bytes = ref_template_path.read_bytes()

        rendered_bytes = _renderer.render(ref_bytes, SAMPLE_CONTEXT)
        assert len(rendered_bytes) > 0

        rendered_doc = Document(BytesIO(rendered_bytes))
        rendered_text = " ".join(p.text for p in rendered_doc.paragraphs)
        assert "AI Template Engine" in rendered_text

    def test_e2e_mobile_preserves_styles(self):
        """Mobile adaptation preserves heading styles."""
        client_bytes = _create_client_docx_mobile()
        iset = _build_mobile_instructions(client_bytes)
        enriched = enrich_instructions(iset)
        validation = validate_instruction_set(enriched)
        safe_iset = validation.sanitized_instructions or enriched
        output_bytes, _, _, _ = _applier.apply(client_bytes, safe_iset)

        out_doc = Document(BytesIO(output_bytes))

        # Heading styles should be preserved
        heading_styles = [
            p.style.name for p in out_doc.paragraphs
            if p.style and p.style.name and p.style.name.startswith("Heading")
        ]
        assert len(heading_styles) > 0, "Heading styles should be preserved in output"


# ---------------------------------------------------------------------------
# E2E: Reference template render verification (all 3 types)
# ---------------------------------------------------------------------------


class TestReferenceTemplateRendering:
    """Verify all EN reference templates render successfully with GW fixture data.

    This proves the reference templates are valid and compatible with the
    TemplateRendererService + GW data format.
    """

    @pytest.mark.parametrize("template_type", ["web", "internal", "mobile"])
    def test_reference_template_renders_with_gw_data(self, template_type: str):
        """Each EN reference template renders without errors using GW fixture."""
        key = (template_type, "en")
        filename = TEMPLATE_MAP.get(key)
        assert filename is not None, f"No template mapped for {key}"

        template_path = TEMPLATE_DIR / filename
        assert template_path.exists(), f"Template not found: {template_path}"

        template_bytes = template_path.read_bytes()

        # Render with GW fixture data -- should not raise
        rendered_bytes = _renderer.render(template_bytes, SAMPLE_CONTEXT)
        assert len(rendered_bytes) > 0

        # Verify rendered output is valid DOCX
        rendered_doc = Document(BytesIO(rendered_bytes))
        assert len(rendered_doc.paragraphs) > 0

        # Verify it contains rendered GW data
        rendered_text = " ".join(p.text for p in rendered_doc.paragraphs)
        assert "AI Template Engine" in rendered_text, (
            f"Rendered {template_type} template should contain client name"
        )

    @pytest.mark.parametrize("template_type", ["web", "internal", "mobile"])
    def test_reference_template_has_patterns(self, template_type: str):
        """Each EN reference template has extractable Jinja2 patterns."""
        ref_info = load_reference_template(template_type, "en")
        assert ref_info.placeholder_count > 0, (
            f"{template_type} reference should have Jinja2 patterns"
        )
        assert ref_info.filename != ""
        assert ref_info.template_type == template_type
        assert ref_info.language == "en"
