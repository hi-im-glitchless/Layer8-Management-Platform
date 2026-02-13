"""Unit tests for DOCX generation service."""
from io import BytesIO

import pytest
from docx import Document
from docx.shared import Pt

from app.services.docx_generator import DocxGeneratorService
from app.services.docx_parser import DocxParserService


@pytest.fixture
def generator():
    """Provide a DocxGeneratorService instance."""
    return DocxGeneratorService()


@pytest.fixture
def parser():
    """Provide a DocxParserService for validating generated output."""
    return DocxParserService()


def _make_template(*paragraphs: str) -> bytes:
    """Helper: create a DOCX template with the given paragraph texts."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_formatted_template() -> bytes:
    """Helper: create a template with formatted runs around a placeholder."""
    doc = Document()
    p = doc.add_paragraph()
    run1 = p.add_run("Name: ")
    run1.bold = True
    run1.font.size = Pt(14)
    run2 = p.add_run("{{ name }}")
    run2.bold = True
    run2.font.size = Pt(14)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestDocxGeneratorSimple:
    """Test simple placeholder replacement."""

    def test_replaces_simple_variable(self, generator):
        template = _make_template("Hello {{ name }}!")
        result = generator.generate(template, {"name": "Alice"})
        doc = Document(BytesIO(result))
        texts = [p.text for p in doc.paragraphs]
        assert any("Alice" in t for t in texts)

    def test_replaces_multiple_variables(self, generator):
        template = _make_template(
            "Client: {{ client }}",
            "Date: {{ date }}",
        )
        result = generator.generate(
            template, {"client": "Acme Corp", "date": "2026-02-13"}
        )
        doc = Document(BytesIO(result))
        texts = [p.text for p in doc.paragraphs]
        assert any("Acme Corp" in t for t in texts)
        assert any("2026-02-13" in t for t in texts)

    def test_handles_missing_context_keys(self, generator):
        """Missing keys should render as empty string (Jinja2 undefined default)."""
        template = _make_template("Hello {{ missing_key }}!")
        result = generator.generate(template, {})
        doc = Document(BytesIO(result))
        # Should not crash; missing key renders as empty or ''
        assert len(doc.paragraphs) >= 1


class TestDocxGeneratorFormatting:
    """Test that formatting is preserved around placeholders."""

    def test_preserves_bold_formatting(self, generator, parser):
        template = _make_formatted_template()
        result = generator.generate(template, {"name": "Bob"})
        structure = parser.parse(result)
        # Find the paragraph that had the placeholder
        formatted = [
            p for p in structure.paragraphs if "Bob" in p.text
        ]
        assert len(formatted) >= 1
        # The runs should still be bold
        bold_runs = [
            r for r in formatted[0].runs if r.bold is True
        ]
        assert len(bold_runs) >= 1


class TestDocxGeneratorOutput:
    """Test that generated output is valid DOCX."""

    def test_output_is_valid_docx(self, generator):
        template = _make_template("Value: {{ x }}")
        result = generator.generate(template, {"x": "42"})
        # Should be parseable as a DOCX
        doc = Document(BytesIO(result))
        assert len(doc.paragraphs) >= 1

    def test_output_can_be_reparsed(self, generator, parser):
        template = _make_template("Score: {{ score }}")
        result = generator.generate(template, {"score": "9.5"})
        structure = parser.parse(result)
        texts = [p.text for p in structure.paragraphs]
        assert any("9.5" in t for t in texts)

    def test_output_bytes_are_nonempty(self, generator):
        template = _make_template("Test")
        result = generator.generate(template, {})
        assert len(result) > 0
        # DOCX files start with PK zip header
        assert result[:2] == b"PK"


class TestDocxGeneratorErrors:
    """Test error handling."""

    def test_rejects_invalid_template_bytes(self, generator):
        with pytest.raises(ValueError, match="Failed to load DOCX template"):
            generator.generate(b"not a docx file", {})

    def test_rejects_empty_template(self, generator):
        with pytest.raises(ValueError, match="Failed to load DOCX template"):
            generator.generate(b"", {})
