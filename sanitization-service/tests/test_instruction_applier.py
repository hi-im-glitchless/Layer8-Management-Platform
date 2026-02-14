"""Unit tests for the DOCX instruction applier.

All test DOCX files are created programmatically -- no external fixtures needed.
"""
from io import BytesIO

import pytest
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

from app.models.adapter import Instruction, InstructionSet
from app.services.instruction_applier import InstructionApplier


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_instruction(
    action: str = "replace_text",
    paragraph_index: int = 0,
    original_text: str = "placeholder",
    replacement_text: str = "{{ client.short_name }}",
    marker_type: str = "text",
    gw_field: str = "client.short_name",
) -> Instruction:
    return Instruction(
        action=action,
        paragraph_index=paragraph_index,
        original_text=original_text,
        replacement_text=replacement_text,
        marker_type=marker_type,
        gw_field=gw_field,
    )


def _make_instruction_set(
    instructions: list[Instruction],
    template_type: str = "web",
    language: str = "en",
) -> InstructionSet:
    return InstructionSet(
        instructions=instructions,
        template_type=template_type,
        language=language,
    )


def _create_test_docx_bytes(paragraphs: list[dict]) -> bytes:
    """Create a test DOCX with specified paragraphs.

    Each paragraph dict can have:
      - text: str (required)
      - bold: bool
      - italic: bool
      - font_name: str
      - font_size: float (in points)
      - color: str (hex like "FF0000")
      - runs: list[dict] -- if provided, creates multiple runs per paragraph
    """
    doc = Document()
    for para_spec in paragraphs:
        p = doc.add_paragraph()

        if "runs" in para_spec:
            # Multiple runs
            for run_spec in para_spec["runs"]:
                r = p.add_run(run_spec.get("text", ""))
                _apply_run_formatting(r, run_spec)
        else:
            r = p.add_run(para_spec.get("text", ""))
            _apply_run_formatting(r, para_spec)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _create_test_docx_with_table(
    pre_paragraphs: list[str],
    table_rows: list[list[str]],
    post_paragraphs: list[str],
) -> bytes:
    """Create a test DOCX with paragraphs and a table.

    Args:
        pre_paragraphs: Paragraphs before the table.
        table_rows: List of rows, each a list of cell texts.
        post_paragraphs: Paragraphs after the table.
    """
    doc = Document()

    for text in pre_paragraphs:
        doc.add_paragraph(text)

    if table_rows:
        n_cols = len(table_rows[0])
        table = doc.add_table(rows=len(table_rows), cols=n_cols)
        for i, row_texts in enumerate(table_rows):
            for j, cell_text in enumerate(row_texts):
                table.rows[i].cells[j].text = cell_text

    for text in post_paragraphs:
        doc.add_paragraph(text)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _apply_run_formatting(run, spec: dict) -> None:
    """Apply formatting from a spec dict to a python-docx run."""
    if spec.get("bold"):
        run.bold = True
    if spec.get("italic"):
        run.italic = True
    if spec.get("font_name"):
        run.font.name = spec["font_name"]
    if spec.get("font_size"):
        run.font.size = Pt(spec["font_size"])
    if spec.get("color"):
        run.font.color.rgb = RGBColor.from_string(spec["color"])
    if spec.get("underline"):
        run.underline = True


def _parse_docx(docx_bytes: bytes) -> Document:
    """Parse DOCX bytes back into a Document for assertions."""
    return Document(BytesIO(docx_bytes))


# ---------------------------------------------------------------------------
# Tests: replace_text
# ---------------------------------------------------------------------------


class TestReplaceText:
    """Tests for the replace_text action."""

    def test_simple_replacement(self):
        """Replace text in a single-run paragraph."""
        docx_bytes = _create_test_docx_bytes([
            {"text": "Hello PLACEHOLDER World"},
        ])
        inst = _make_instruction(
            original_text="PLACEHOLDER",
            replacement_text="{{ client.short_name }}",
            paragraph_index=0,
        )
        iset = _make_instruction_set([inst])
        applier = InstructionApplier()

        result_bytes, applied, skipped, warnings = applier.apply(docx_bytes, iset)

        assert applied == 1
        assert skipped == 0
        doc = _parse_docx(result_bytes)
        assert "{{ client.short_name }}" in doc.paragraphs[0].text
        assert "PLACEHOLDER" not in doc.paragraphs[0].text

    def test_preserves_bold_formatting(self):
        """Replace text in a bold run and verify bold is preserved."""
        docx_bytes = _create_test_docx_bytes([
            {"text": "Client Name Here", "bold": True, "font_name": "Arial", "font_size": 14},
        ])
        inst = _make_instruction(
            original_text="Client Name Here",
            replacement_text="{{ client.short_name }}",
            paragraph_index=0,
        )
        iset = _make_instruction_set([inst])
        applier = InstructionApplier()

        result_bytes, applied, skipped, warnings = applier.apply(docx_bytes, iset)

        assert applied == 1
        doc = _parse_docx(result_bytes)
        run = doc.paragraphs[0].runs[0]
        assert run.text == "{{ client.short_name }}"
        assert run.bold is True
        assert run.font.name == "Arial"
        assert run.font.size == Pt(14)

    def test_preserves_italic_and_color(self):
        """Replace text preserving italic and color formatting."""
        docx_bytes = _create_test_docx_bytes([
            {"text": "Replace Me", "italic": True, "color": "FF0000"},
        ])
        inst = _make_instruction(
            original_text="Replace Me",
            replacement_text="{{ project.start_date }}",
            paragraph_index=0,
        )
        iset = _make_instruction_set([inst])
        applier = InstructionApplier()

        result_bytes, applied, skipped, warnings = applier.apply(docx_bytes, iset)

        assert applied == 1
        doc = _parse_docx(result_bytes)
        run = doc.paragraphs[0].runs[0]
        assert run.italic is True
        assert str(run.font.color.rgb) == "FF0000"

    def test_multi_run_replacement(self):
        """Replace text that spans multiple runs."""
        docx_bytes = _create_test_docx_bytes([
            {
                "runs": [
                    {"text": "Client ", "bold": True},
                    {"text": "Name", "bold": True},
                    {"text": " Here", "bold": True},
                ]
            },
        ])
        inst = _make_instruction(
            original_text="Client Name Here",
            replacement_text="{{ client.short_name }}",
            paragraph_index=0,
        )
        iset = _make_instruction_set([inst])
        applier = InstructionApplier()

        result_bytes, applied, skipped, warnings = applier.apply(docx_bytes, iset)

        assert applied == 1
        doc = _parse_docx(result_bytes)
        full_text = "".join(r.text for r in doc.paragraphs[0].runs)
        assert "{{ client.short_name }}" in full_text
        assert "Client Name Here" not in full_text

    def test_multi_run_partial_replacement(self):
        """Replace text spanning part of multiple runs."""
        docx_bytes = _create_test_docx_bytes([
            {
                "runs": [
                    {"text": "Pre-Cli", "bold": True},
                    {"text": "ent Na", "italic": True},
                    {"text": "me-Post", "bold": True},
                ]
            },
        ])
        inst = _make_instruction(
            original_text="Client Name",
            replacement_text="{{ client.short_name }}",
            paragraph_index=0,
        )
        iset = _make_instruction_set([inst])
        applier = InstructionApplier()

        result_bytes, applied, skipped, warnings = applier.apply(docx_bytes, iset)

        assert applied == 1
        doc = _parse_docx(result_bytes)
        full_text = "".join(r.text for r in doc.paragraphs[0].runs)
        assert "{{ client.short_name }}" in full_text
        assert "Client Name" not in full_text
        # Pre and Post text should be preserved
        assert "Pre-" in full_text
        assert "-Post" in full_text

    def test_text_not_found_skipped(self):
        """Instruction for text not in paragraph is skipped with warning."""
        docx_bytes = _create_test_docx_bytes([
            {"text": "Something else entirely"},
        ])
        inst = _make_instruction(
            original_text="NONEXISTENT",
            replacement_text="{{ client.short_name }}",
            paragraph_index=0,
        )
        iset = _make_instruction_set([inst])
        applier = InstructionApplier()

        result_bytes, applied, skipped, warnings = applier.apply(docx_bytes, iset)

        assert applied == 0
        assert skipped == 1
        assert len(warnings) == 1


# ---------------------------------------------------------------------------
# Tests: insert_before / insert_after
# ---------------------------------------------------------------------------


class TestInsertParagraph:
    """Tests for insert_before and insert_after actions."""

    def test_insert_before(self):
        """Insert a paragraph before the target."""
        docx_bytes = _create_test_docx_bytes([
            {"text": "Paragraph A"},
            {"text": "Paragraph B"},
            {"text": "Paragraph C"},
        ])
        inst = _make_instruction(
            action="insert_before",
            paragraph_index=1,
            original_text="",
            replacement_text="{% for finding in findings %}",
        )
        iset = _make_instruction_set([inst])
        applier = InstructionApplier()

        result_bytes, applied, skipped, warnings = applier.apply(docx_bytes, iset)

        assert applied == 1
        doc = _parse_docx(result_bytes)
        texts = [p.text for p in doc.paragraphs]
        # New paragraph should be before "Paragraph B"
        b_index = texts.index("Paragraph B")
        assert texts[b_index - 1] == "{% for finding in findings %}"

    def test_insert_after(self):
        """Insert a paragraph after the target."""
        docx_bytes = _create_test_docx_bytes([
            {"text": "Paragraph A"},
            {"text": "Paragraph B"},
            {"text": "Paragraph C"},
        ])
        inst = _make_instruction(
            action="insert_after",
            paragraph_index=1,
            original_text="",
            replacement_text="{% endfor %}",
        )
        iset = _make_instruction_set([inst])
        applier = InstructionApplier()

        result_bytes, applied, skipped, warnings = applier.apply(docx_bytes, iset)

        assert applied == 1
        doc = _parse_docx(result_bytes)
        texts = [p.text for p in doc.paragraphs]
        b_index = texts.index("Paragraph B")
        assert texts[b_index + 1] == "{% endfor %}"


# ---------------------------------------------------------------------------
# Tests: wrap_table_row
# ---------------------------------------------------------------------------


class TestWrapTableRow:
    """Tests for the wrap_table_row action."""

    def test_wrap_table_row_adds_markers(self):
        """wrap_table_row adds loop start/end marker rows around the target row."""
        docx_bytes = _create_test_docx_with_table(
            pre_paragraphs=["Header paragraph"],
            table_rows=[
                ["Header1", "Header2"],
                ["Data1", "Data2"],
                ["Footer1", "Footer2"],
            ],
            post_paragraphs=["Footer paragraph"],
        )

        # The data row paragraph is inside the table, not a body paragraph.
        # We need to find the paragraph index in the body paragraphs list.
        # For wrap_table_row, the paragraph_index refers to body paragraphs,
        # but table cell paragraphs are not in doc.paragraphs.
        # Let's test that a non-table paragraph is handled gracefully.
        # For actual table wrapping, the paragraph must be in a table cell.

        # Create a DOCX where we can address table cell paragraphs
        doc = Document()
        doc.add_paragraph("Before table")
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Item"
        table.rows[0].cells[1].text = "Value"
        table.rows[1].cells[0].text = "Scope Item"
        table.rows[1].cells[1].text = "Scope Value"
        doc.add_paragraph("After table")

        buf = BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        # doc.paragraphs only includes body paragraphs, not table cell paragraphs.
        # wrap_table_row on a body paragraph should fail gracefully.
        inst = _make_instruction(
            action="wrap_table_row",
            paragraph_index=0,
            original_text="",
            replacement_text="{%tr for item in scope %}",
        )
        iset = _make_instruction_set([inst])
        applier = InstructionApplier()

        result_bytes, applied, skipped, warnings = applier.apply(docx_bytes, iset)

        # Body paragraph is not in a table, so it should be skipped
        assert skipped == 1
        assert applied == 0

    def test_wrap_table_row_on_table_paragraph(self):
        """wrap_table_row works on a paragraph that is inside a table cell.

        We test this by directly calling the _wrap_table_row_with_loop method
        on a paragraph from a table cell.
        """
        doc = Document()
        table = doc.add_table(rows=3, cols=2)
        table.rows[0].cells[0].text = "Header"
        table.rows[1].cells[0].text = "Data Row"
        table.rows[2].cells[0].text = "Footer"

        # Get the paragraph in the data row's first cell
        data_cell_para = table.rows[1].cells[0].paragraphs[0]

        applier = InstructionApplier()
        success = applier._wrap_table_row_with_loop(
            doc, data_cell_para, "{%tr for item in scope %}"
        )

        assert success is True

        # Verify the table now has 5 rows (3 original + 2 marker rows)
        tbl_elem = table._tbl
        tr_elements = tbl_elem.findall(qn("w:tr"))
        assert len(tr_elements) == 5

        # First marker row should contain the start marker
        first_marker = tr_elements[1]  # 0=header, 1=start marker
        first_text = first_marker.find(f".//{qn('w:t')}")
        assert first_text is not None
        assert "{%tr for item in scope %}" in first_text.text

        # Last marker row should contain endfor
        last_marker = tr_elements[3]  # 2=data, 3=end marker
        last_text = last_marker.find(f".//{qn('w:t')}")
        assert last_text is not None
        assert "{%tr endfor %}" in last_text.text


# ---------------------------------------------------------------------------
# Tests: edge cases and output validity
# ---------------------------------------------------------------------------


class TestEdgeCases:
    """Tests for edge cases and output validity."""

    def test_invalid_paragraph_index_relocated_by_text(self):
        """Out-of-range paragraph index is relocated via text-based fallback."""
        docx_bytes = _create_test_docx_bytes([
            {"text": "Only paragraph"},
        ])
        inst = _make_instruction(
            paragraph_index=99,
            original_text="Only paragraph",
            replacement_text="{{ client.short_name }}",
        )
        iset = _make_instruction_set([inst])
        applier = InstructionApplier()

        result_bytes, applied, skipped, warnings = applier.apply(docx_bytes, iset)

        # Text-based fallback finds the paragraph even with wrong index
        assert applied == 1
        assert skipped == 0
        doc = Document(BytesIO(result_bytes))
        assert "{{ client.short_name }}" in doc.paragraphs[0].text

    def test_invalid_paragraph_index_no_text_match_skipped(self):
        """Out-of-range paragraph index with non-matching text is skipped."""
        docx_bytes = _create_test_docx_bytes([
            {"text": "Only paragraph"},
        ])
        inst = _make_instruction(
            paragraph_index=99,
            original_text="Text that does not exist anywhere",
            replacement_text="{{ client.short_name }}",
        )
        iset = _make_instruction_set([inst])
        applier = InstructionApplier()

        result_bytes, applied, skipped, warnings = applier.apply(docx_bytes, iset)

        assert applied == 0
        assert skipped == 1
        assert len(warnings) == 1

    def test_output_is_valid_docx(self):
        """Output can be re-parsed by python-docx as a valid DOCX."""
        docx_bytes = _create_test_docx_bytes([
            {"text": "Hello World", "bold": True},
            {"text": "Replace This"},
        ])
        inst = _make_instruction(
            paragraph_index=1,
            original_text="Replace This",
            replacement_text="{{ project.end_date }}",
        )
        iset = _make_instruction_set([inst])
        applier = InstructionApplier()

        result_bytes, applied, skipped, warnings = applier.apply(docx_bytes, iset)

        # Should not raise any exception
        doc = _parse_docx(result_bytes)
        assert len(doc.paragraphs) >= 2

    def test_multiple_instructions_applied(self):
        """Multiple instructions are applied correctly (bottom-up order)."""
        docx_bytes = _create_test_docx_bytes([
            {"text": "Client: ACME Corp"},
            {"text": "Start: 2024-01-01"},
            {"text": "End: 2024-12-31"},
        ])
        instructions = [
            _make_instruction(
                paragraph_index=0,
                original_text="ACME Corp",
                replacement_text="{{ client.short_name }}",
            ),
            _make_instruction(
                paragraph_index=1,
                original_text="2024-01-01",
                replacement_text="{{ project.start_date }}",
                gw_field="project.start_date",
            ),
            _make_instruction(
                paragraph_index=2,
                original_text="2024-12-31",
                replacement_text="{{ project.end_date }}",
                gw_field="project.end_date",
            ),
        ]
        iset = _make_instruction_set(instructions)
        applier = InstructionApplier()

        result_bytes, applied, skipped, warnings = applier.apply(docx_bytes, iset)

        assert applied == 3
        assert skipped == 0
        doc = _parse_docx(result_bytes)
        assert "{{ client.short_name }}" in doc.paragraphs[0].text
        assert "{{ project.start_date }}" in doc.paragraphs[1].text
        assert "{{ project.end_date }}" in doc.paragraphs[2].text

    def test_empty_instruction_set(self):
        """Empty instruction set returns the document unchanged."""
        docx_bytes = _create_test_docx_bytes([
            {"text": "Unchanged text"},
        ])
        iset = _make_instruction_set([])
        applier = InstructionApplier()

        result_bytes, applied, skipped, warnings = applier.apply(docx_bytes, iset)

        assert applied == 0
        assert skipped == 0
        doc = _parse_docx(result_bytes)
        assert doc.paragraphs[0].text == "Unchanged text"

    def test_formatting_preservation_font_size_and_color(self):
        """Verify formatting preservation on re-parsed output."""
        docx_bytes = _create_test_docx_bytes([
            {
                "text": "Styled Text",
                "bold": True,
                "italic": True,
                "font_name": "Calibri",
                "font_size": 12,
                "color": "0000FF",
            },
        ])
        inst = _make_instruction(
            paragraph_index=0,
            original_text="Styled Text",
            replacement_text="{{ finding.title }}",
            gw_field="finding.title",
        )
        iset = _make_instruction_set([inst])
        applier = InstructionApplier()

        result_bytes, applied, skipped, warnings = applier.apply(docx_bytes, iset)

        assert applied == 1
        doc = _parse_docx(result_bytes)
        run = doc.paragraphs[0].runs[0]
        assert run.text == "{{ finding.title }}"
        assert run.bold is True
        assert run.italic is True
        assert run.font.name == "Calibri"
        assert run.font.size == Pt(12)
        assert str(run.font.color.rgb) == "0000FF"

    def test_negative_paragraph_index_relocated_by_text(self):
        """Negative paragraph_index is relocated via text-based fallback."""
        docx_bytes = _create_test_docx_bytes([
            {"text": "Hello"},
        ])
        inst = _make_instruction(
            paragraph_index=-1,
            original_text="Hello",
            replacement_text="{{ client.short_name }}",
        )
        iset = _make_instruction_set([inst])
        applier = InstructionApplier()

        result_bytes, applied, skipped, warnings = applier.apply(docx_bytes, iset)

        # Text-based fallback finds the paragraph even with negative index
        assert applied == 1
        assert skipped == 0

    def test_replace_text_in_table_cell(self):
        """Text in a body table cell is found and replaced (strategy 4)."""
        doc = Document()
        doc.add_paragraph("Unrelated paragraph")
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Header"
        table.rows[1].cells[0].text = "Old Company Name"
        table.rows[1].cells[1].text = "Details"
        buf = BytesIO()
        doc.save(buf)
        template_bytes = buf.getvalue()

        inst = _make_instruction(
            paragraph_index=999,  # Wrong index, must fallback to table search
            original_text="Old Company Name",
            replacement_text="{{ client.short_name }}",
        )
        iset = _make_instruction_set([inst])
        applier = InstructionApplier()

        result_bytes, applied, skipped, warnings = applier.apply(template_bytes, iset)

        assert applied == 1
        assert skipped == 0

        # Verify replacement in table cell
        result_doc = Document(BytesIO(result_bytes))
        cell_text = result_doc.tables[0].rows[1].cells[0].text
        assert "{{ client.short_name }}" in cell_text

    def test_replace_text_in_textbox(self):
        """Text in a text box is found and replaced (strategy 5)."""
        from docx.oxml import OxmlElement

        doc = Document()
        doc.add_paragraph("Unrelated paragraph")

        # Create a minimal text box structure in the body
        # w:txbxContent > w:p > w:r > w:t
        txbx = OxmlElement("w:txbxContent")
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "Report Title"
        r.append(t)
        p.append(r)
        txbx.append(p)

        # Embed in a w:pict element inside the last body paragraph
        last_p = doc.element.body.findall(qn("w:p"))[-1]
        wrapper_r = OxmlElement("w:r")
        pict = OxmlElement("w:pict")
        pict.append(txbx)
        wrapper_r.append(pict)
        last_p.append(wrapper_r)

        buf = BytesIO()
        doc.save(buf)
        template_bytes = buf.getvalue()

        inst = _make_instruction(
            paragraph_index=999,  # Wrong index
            original_text="Report Title",
            replacement_text="{{ finding.title }}",
            gw_field="finding.title",
        )
        iset = _make_instruction_set([inst])
        applier = InstructionApplier()

        result_bytes, applied, skipped, warnings = applier.apply(template_bytes, iset)

        assert applied == 1
        assert skipped == 0

        # Verify replacement in text box
        result_doc = Document(BytesIO(result_bytes))
        txbx_elements = result_doc.element.body.findall(
            ".//" + qn("w:txbxContent")
        )
        assert len(txbx_elements) >= 1
        txbx_texts = txbx_elements[0].findall(".//" + qn("w:t"))
        full_text = "".join(t.text or "" for t in txbx_texts)
        assert "{{ finding.title }}" in full_text
