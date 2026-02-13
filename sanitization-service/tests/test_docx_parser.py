"""Unit tests for DOCX parsing service."""
import os
from io import BytesIO

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.services.docx_parser import DocxParserService

FIXTURES_DIR = os.path.join(os.path.dirname(__file__), "fixtures")


@pytest.fixture
def parser():
    """Provide a DocxParserService instance."""
    return DocxParserService()


@pytest.fixture
def sample_docx_bytes():
    """Load the sample.docx fixture as bytes."""
    path = os.path.join(FIXTURES_DIR, "sample.docx")
    with open(path, "rb") as f:
        return f.read()


@pytest.fixture
def minimal_docx_bytes():
    """Create a minimal DOCX in memory with known content."""
    doc = Document()
    doc.core_properties.author = "Test Author"
    doc.core_properties.title = "Test Title"
    doc.add_heading("Main Heading", level=1)
    doc.add_heading("Sub Heading", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Formatted text")
    run.bold = True
    run.italic = True
    run.font.name = "Courier New"
    run.font.size = Pt(14)

    p2 = doc.add_paragraph("Centered")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def empty_docx_bytes():
    """Create an empty DOCX (no user content)."""
    doc = Document()
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestDocxParserParagraphs:
    """Test paragraph extraction."""

    def test_extracts_paragraphs_with_text(self, parser, sample_docx_bytes):
        result = parser.parse(sample_docx_bytes)
        texts = [p.text for p in result.paragraphs]
        assert any("Bold text" in t for t in texts)

    def test_extracts_heading_levels(self, parser, minimal_docx_bytes):
        result = parser.parse(minimal_docx_bytes)
        headings = [p for p in result.paragraphs if p.heading_level is not None]
        levels = {p.heading_level for p in headings}
        assert 1 in levels
        assert 2 in levels

    def test_extracts_style_names(self, parser, minimal_docx_bytes):
        result = parser.parse(minimal_docx_bytes)
        styles = [p.style_name for p in result.paragraphs]
        assert "Heading 1" in styles
        assert "Heading 2" in styles

    def test_extracts_paragraph_alignment(self, parser, minimal_docx_bytes):
        result = parser.parse(minimal_docx_bytes)
        centered = [p for p in result.paragraphs if p.alignment == "CENTER"]
        assert len(centered) >= 1
        assert centered[0].text == "Centered"

    def test_extracts_run_formatting(self, parser, minimal_docx_bytes):
        result = parser.parse(minimal_docx_bytes)
        # Find the paragraph with formatted runs
        formatted = [p for p in result.paragraphs if p.text == "Formatted text"]
        assert len(formatted) == 1
        runs = formatted[0].runs
        assert len(runs) == 1
        assert runs[0].bold is True
        assert runs[0].italic is True
        assert runs[0].font_name == "Courier New"
        assert runs[0].font_size == 14.0


class TestDocxParserTables:
    """Test table extraction."""

    def test_extracts_tables(self, parser, sample_docx_bytes):
        result = parser.parse(sample_docx_bytes)
        assert len(result.tables) >= 1

    def test_table_has_correct_rows_and_cells(self, parser, minimal_docx_bytes):
        result = parser.parse(minimal_docx_bytes)
        assert len(result.tables) == 1
        table = result.tables[0]
        assert len(table.rows) == 2
        assert len(table.rows[0].cells) == 2
        assert table.rows[0].cells[0].text == "A"
        assert table.rows[0].cells[1].text == "B"
        assert table.rows[1].cells[0].text == "C"
        assert table.rows[1].cells[1].text == "D"

    def test_table_cells_have_nested_paragraphs(self, parser, minimal_docx_bytes):
        result = parser.parse(minimal_docx_bytes)
        cell = result.tables[0].rows[0].cells[0]
        assert len(cell.paragraphs) >= 1
        assert cell.paragraphs[0].text == "A"

    def test_sample_table_content(self, parser, sample_docx_bytes):
        result = parser.parse(sample_docx_bytes)
        table = result.tables[0]
        assert table.rows[0].cells[0].text == "ID"
        assert table.rows[1].cells[1].text == "XSS Reflected"
        assert table.rows[2].cells[2].text == "Critical"


class TestDocxParserSectionsAndMeta:
    """Test section and metadata extraction."""

    def test_extracts_sections(self, parser, sample_docx_bytes):
        result = parser.parse(sample_docx_bytes)
        assert len(result.sections) >= 1
        section = result.sections[0]
        assert section.page_width is not None
        assert section.page_height is not None

    def test_extracts_metadata(self, parser, sample_docx_bytes):
        result = parser.parse(sample_docx_bytes)
        assert result.metadata.get("author") == "Layer8 Test"
        assert result.metadata.get("title") == "Sample Test Document"

    def test_extracts_unique_styles(self, parser, sample_docx_bytes):
        result = parser.parse(sample_docx_bytes)
        assert len(result.styles) > 0
        assert "Heading 1" in result.styles


class TestDocxParserEdgeCases:
    """Test edge cases and error handling."""

    def test_empty_document(self, parser, empty_docx_bytes):
        result = parser.parse(empty_docx_bytes)
        # Empty doc still has a default section
        assert len(result.sections) >= 1
        # Should not crash
        assert isinstance(result.paragraphs, list)
        assert isinstance(result.tables, list)

    def test_rejects_non_docx_bytes(self, parser):
        with pytest.raises(ValueError, match="Failed to parse DOCX"):
            parser.parse(b"This is not a DOCX file")

    def test_rejects_empty_bytes(self, parser):
        with pytest.raises(ValueError, match="Failed to parse DOCX"):
            parser.parse(b"")
