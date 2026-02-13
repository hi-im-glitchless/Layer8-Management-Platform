"""Integration tests for adapter API endpoints.

Tests the FastAPI routes (/adapter/analyze, /validate-mapping, /apply, /enrich)
with realistic request payloads. Uses FastAPI TestClient (synchronous wrapper
around httpx.AsyncClient).

No external services required -- all tests use programmatically created DOCX files.
"""
import base64
import json
from io import BytesIO

import pytest
from docx import Document as DocxDoc
from fastapi.testclient import TestClient

from app.main import app


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def client():
    """FastAPI test client."""
    return TestClient(app)


@pytest.fixture
def test_docx_base64() -> str:
    """Base64-encoded minimal DOCX for testing."""
    doc = DocxDoc()
    doc.add_heading("Security Assessment Report", level=1)
    doc.add_paragraph("Client Name: Acme Corp")
    doc.add_paragraph("Assessment Period: January 2025 - February 2025")
    doc.add_heading("Findings", level=2)
    doc.add_paragraph("1. SQL Injection vulnerability in login endpoint.")
    doc.add_paragraph("Scope: www.example.com")
    buf = BytesIO()
    doc.save(buf)
    return base64.b64encode(buf.getvalue()).decode("ascii")


@pytest.fixture
def valid_instruction_set() -> dict:
    """A valid instruction set JSON for the /apply endpoint."""
    return {
        "instructions": [
            {
                "action": "replace_text",
                "paragraph_index": 1,
                "original_text": "Acme Corp",
                "replacement_text": "{{ client.short_name }}",
                "marker_type": "text",
                "gw_field": "client.short_name",
            },
            {
                "action": "replace_text",
                "paragraph_index": 2,
                "original_text": "January 2025 - February 2025",
                "replacement_text": "{{ project.start_date }} - {{ project.end_date }}",
                "marker_type": "text",
                "gw_field": "project.start_date",
            },
        ],
        "template_type": "web",
        "language": "en",
    }


# ---------------------------------------------------------------------------
# Tests: POST /adapter/analyze
# ---------------------------------------------------------------------------


class TestAnalyzeEndpointIntegration:
    """Integration tests for POST /adapter/analyze."""

    def test_analyze_endpoint_returns_prompt_and_reference_info(
        self, client, test_docx_base64
    ):
        """Verify analyze returns non-empty prompt and valid SHA-256 hash."""
        response = client.post(
            "/adapter/analyze",
            json={
                "template_base64": test_docx_base64,
                "template_type": "web",
                "language": "en",
            },
        )
        assert response.status_code == 200
        data = response.json()

        # Prompt should be non-empty
        assert isinstance(data["prompt"], str)
        assert len(data["prompt"]) > 100

        # System prompt should be non-empty
        assert isinstance(data["system_prompt"], str)
        assert len(data["system_prompt"]) > 0

        # Reference template hash should be a valid hex SHA-256 (64 chars)
        assert isinstance(data["reference_template_hash"], str)
        assert len(data["reference_template_hash"]) == 64
        # Should be valid hex
        int(data["reference_template_hash"], 16)

        # Paragraph count should be positive
        assert data["paragraph_count"] > 0

    def test_analyze_rejects_non_docx(self, client):
        """POST /adapter/analyze with plain text file returns 400."""
        fake_b64 = base64.b64encode(b"This is not a DOCX file at all").decode("ascii")
        response = client.post(
            "/adapter/analyze",
            json={
                "template_base64": fake_b64,
                "template_type": "web",
                "language": "en",
            },
        )
        assert response.status_code == 400
        assert "not a valid DOCX" in response.json()["detail"]


# ---------------------------------------------------------------------------
# Tests: POST /adapter/validate-mapping
# ---------------------------------------------------------------------------


class TestValidateMappingEndpointIntegration:
    """Integration tests for POST /adapter/validate-mapping."""

    def test_validate_mapping_accepts_valid_plan(self, client):
        """Valid mapping plan returns valid=True and parsed MappingPlan."""
        llm_response = json.dumps({
            "entries": [
                {
                    "section_index": 0,
                    "section_text": "Client Name: Acme Corp",
                    "gw_field": "client.short_name",
                    "placeholder_template": "{{ client.short_name }}",
                    "confidence": 0.95,
                    "marker_type": "text",
                    "rationale": "Contains client name pattern",
                },
                {
                    "section_index": 2,
                    "section_text": "Assessment Period",
                    "gw_field": "project.start_date",
                    "placeholder_template": "{{ project.start_date }}",
                    "confidence": 0.85,
                    "marker_type": "text",
                    "rationale": "Contains date range",
                },
            ],
            "warnings": ["Some headings were ambiguous"],
        })
        response = client.post(
            "/adapter/validate-mapping",
            json={
                "llm_response": llm_response,
                "template_type": "web",
                "language": "en",
                "paragraph_count": 10,
            },
        )
        assert response.status_code == 200
        data = response.json()
        assert data["valid"] is True
        assert data["mapping_plan"] is not None
        assert len(data["mapping_plan"]["entries"]) == 2

    def test_validate_mapping_rejects_invalid_index(self, client):
        """Section index 999 (beyond doc length) returns valid=False."""
        llm_response = json.dumps({
            "entries": [
                {
                    "section_index": 999,
                    "section_text": "out of range section",
                    "gw_field": "client.short_name",
                    "placeholder_template": "{{ client.short_name }}",
                    "confidence": 0.9,
                    "marker_type": "text",
                    "rationale": "test",
                }
            ],
            "warnings": [],
        })
        response = client.post(
            "/adapter/validate-mapping",
            json={
                "llm_response": llm_response,
                "template_type": "web",
                "language": "en",
                "paragraph_count": 10,
            },
        )
        data = response.json()
        assert data["valid"] is False
        assert len(data["errors"]) > 0
        assert any("out of range" in e for e in data["errors"])


# ---------------------------------------------------------------------------
# Tests: POST /adapter/apply
# ---------------------------------------------------------------------------


class TestApplyEndpointIntegration:
    """Integration tests for POST /adapter/apply."""

    def test_apply_endpoint_returns_modified_docx(
        self, client, test_docx_base64, valid_instruction_set
    ):
        """Apply endpoint returns output_base64 containing a valid DOCX."""
        response = client.post(
            "/adapter/apply",
            json={
                "template_base64": test_docx_base64,
                "instruction_set": valid_instruction_set,
            },
        )
        assert response.status_code == 200
        data = response.json()

        # output_base64 should be non-empty
        assert isinstance(data["output_base64"], str)
        assert len(data["output_base64"]) > 0

        # Decode output and verify it is valid DOCX (PK magic bytes)
        output_bytes = base64.b64decode(data["output_base64"])
        assert output_bytes[:4] == b"PK\x03\x04", "Output should be a valid ZIP/DOCX file"

        # applied_count should be > 0
        assert data["applied_count"] > 0

        # Verify the output is parseable by python-docx
        doc = DocxDoc(BytesIO(output_bytes))
        assert len(doc.paragraphs) > 0

        # Verify that the replacement was applied
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "{{ client.short_name }}" in all_text

    def test_apply_rejects_unsafe_jinja2(self, client, test_docx_base64):
        """Apply with unsafe Jinja2 instructions returns validation error."""
        unsafe_instruction_set = {
            "instructions": [
                {
                    "action": "replace_text",
                    "paragraph_index": 1,
                    "original_text": "Acme Corp",
                    "replacement_text": "{{ import('os').system('rm -rf /') }}",
                    "marker_type": "text",
                    "gw_field": "client.short_name",
                },
            ],
            "template_type": "web",
            "language": "en",
        }
        response = client.post(
            "/adapter/apply",
            json={
                "template_base64": test_docx_base64,
                "instruction_set": unsafe_instruction_set,
            },
        )
        # Should return 422 because all instructions are invalid
        assert response.status_code == 422
        data = response.json()
        assert "detail" in data


# ---------------------------------------------------------------------------
# Tests: POST /adapter/enrich
# ---------------------------------------------------------------------------


class TestEnrichEndpointIntegration:
    """Integration tests for POST /adapter/enrich."""

    def test_enrich_endpoint_adds_markers(self, client):
        """Enrich endpoint applies marker rules and type features."""
        instruction_set = {
            "instructions": [
                {
                    "action": "replace_text",
                    "paragraph_index": 0,
                    "original_text": "Client Name",
                    "replacement_text": "{{ client.short_name }}",
                    "marker_type": "text",
                    "gw_field": "client.short_name",
                },
                {
                    "action": "replace_text",
                    "paragraph_index": 1,
                    "original_text": "Finding description here",
                    "replacement_text": "{{ finding.description_rt }}",
                    "marker_type": "paragraph_rt",
                    "gw_field": "finding.description_rt",
                },
            ],
            "template_type": "web",
            "language": "en",
        }
        response = client.post(
            "/adapter/enrich",
            json={"instruction_set": instruction_set},
        )
        assert response.status_code == 200
        data = response.json()

        enriched_iset = data["instruction_set"]
        assert len(enriched_iset["instructions"]) == 2

        # The description_rt field should have been rewritten to {{p ... }}
        desc_inst = enriched_iset["instructions"][1]
        assert "{{p" in desc_inst["replacement_text"], (
            f"Expected {{{{p marker for paragraph_rt field, got: {desc_inst['replacement_text']}"
        )

        # Web template should have additional_blocks for scope loops
        assert len(enriched_iset["additional_blocks"]) > 0

    def test_enrich_internal_template_adds_filter_type(self, client):
        """Enrich for internal template adds filter_type blocks."""
        instruction_set = {
            "instructions": [
                {
                    "action": "replace_text",
                    "paragraph_index": 0,
                    "original_text": "Client Name",
                    "replacement_text": "{{ client.short_name }}",
                    "marker_type": "text",
                    "gw_field": "client.short_name",
                },
            ],
            "template_type": "internal",
            "language": "en",
        }
        response = client.post(
            "/adapter/enrich",
            json={"instruction_set": instruction_set},
        )
        assert response.status_code == 200
        data = response.json()

        additional = " ".join(data["instruction_set"]["additional_blocks"])
        assert "filter_type" in additional
        assert "namespace" in additional
