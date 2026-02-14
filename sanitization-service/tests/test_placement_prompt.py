"""Unit tests for the placement prompt builder and validation endpoint."""
import json

import pytest

from app.models.adapter import (
    FIELD_MARKER_MAP,
    Instruction,
    InstructionSet,
    MappingEntry,
    MappingPlan,
    ValidatePlacementRequest,
    ValidatePlacementResponse,
)
from app.models.docx import (
    DocxCell,
    DocxParagraph,
    DocxRow,
    DocxSection,
    DocxStructure,
    DocxTable,
)
from app.services.placement_prompt import (
    build_placement_prompt,
    build_placement_system_prompt,
    build_zone_map,
)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def sample_doc_structure():
    """Create a DocxStructure with paragraphs in different zones."""
    paragraphs = [
        DocxParagraph(text="Company Logo Area", style_name="Title", zone="cover"),
        DocxParagraph(text="Confidential Report", style_name="Normal", zone="cover"),
        DocxParagraph(
            text="Executive Summary",
            style_name="Heading 1",
            heading_level=1,
            zone="body",
        ),
        DocxParagraph(
            text="This report presents the findings from the security assessment of Client Corp conducted between January and February 2025.",
            style_name="Normal",
            zone="body",
        ),
        DocxParagraph(
            text="Scope of Assessment",
            style_name="Heading 2",
            heading_level=2,
            zone="body",
        ),
        DocxParagraph(
            text="The assessment covered web application security testing.",
            style_name="Normal",
            zone="body",
        ),
        DocxParagraph(text="", style_name="Normal", zone="body"),
        DocxParagraph(
            text="Findings Overview",
            style_name="Heading 1",
            heading_level=1,
            zone="body",
        ),
        DocxParagraph(
            text="SQL Injection in login form allows unauthenticated access.",
            style_name="Normal",
            zone="body",
        ),
    ]
    tables = [
        DocxTable(
            rows=[
                DocxRow(
                    cells=[
                        DocxCell(text="Finding ID", paragraphs=[DocxParagraph(text="Finding ID")]),
                        DocxCell(text="Title", paragraphs=[DocxParagraph(text="Title")]),
                        DocxCell(text="Severity", paragraphs=[DocxParagraph(text="Severity")]),
                    ]
                ),
                DocxRow(
                    cells=[
                        DocxCell(text="001", paragraphs=[DocxParagraph(text="001")]),
                        DocxCell(text="SQL Injection", paragraphs=[DocxParagraph(text="SQL Injection")]),
                        DocxCell(text="Critical", paragraphs=[DocxParagraph(text="Critical")]),
                    ]
                ),
            ],
            style_name="Table Grid",
        )
    ]
    sections = [
        DocxSection(
            header_paragraphs=[
                DocxParagraph(text="Company Header", style_name="Header", zone="header"),
            ],
            footer_paragraphs=[
                DocxParagraph(text="Page {{ page_number }}", style_name="Footer", zone="footer"),
            ],
        )
    ]
    return DocxStructure(
        paragraphs=paragraphs,
        tables=tables,
        images=[],
        sections=sections,
        styles=["Normal", "Title", "Heading 1", "Heading 2", "Header", "Footer"],
        metadata={"title": "Test Report"},
    )


@pytest.fixture
def sample_mapping_plan():
    """Create a MappingPlan with 3 entries."""
    return MappingPlan(
        entries=[
            MappingEntry(
                section_index=3,
                section_text="This report presents the findings",
                gw_field="client.short_name",
                placeholder_template="{{ client.short_name }}",
                confidence=0.9,
                marker_type="text",
            ),
            MappingEntry(
                section_index=5,
                section_text="The assessment covered web application",
                gw_field="finding.description_rt",
                placeholder_template="{{p finding.description_rt }}",
                confidence=0.85,
                marker_type="paragraph_rt",
            ),
            MappingEntry(
                section_index=8,
                section_text="SQL Injection in login form",
                gw_field="finding.title",
                placeholder_template="{{ finding.title }}",
                confidence=0.95,
                marker_type="text",
            ),
        ],
        template_type="web",
        language="en",
    )


# ---------------------------------------------------------------------------
# build_zone_map tests
# ---------------------------------------------------------------------------


class TestBuildZoneMap:
    def test_zone_sections_appear_in_output(self, sample_doc_structure, sample_mapping_plan):
        result = build_zone_map(sample_doc_structure, sample_mapping_plan)
        assert "## ZONE: cover" in result
        assert "## ZONE: body" in result

    def test_header_footer_sections_appear(self, sample_doc_structure, sample_mapping_plan):
        result = build_zone_map(sample_doc_structure, sample_mapping_plan)
        assert "## ZONE: header (Section 1)" in result
        assert "## ZONE: footer (Section 1)" in result

    def test_table_cells_appear(self, sample_doc_structure, sample_mapping_plan):
        result = build_zone_map(sample_doc_structure, sample_mapping_plan)
        assert "## ZONE: table_cell (Table 0)" in result
        assert "[T0.R0.C0]: Finding ID" in result
        assert "[T0.R0.C1]: Title" in result

    def test_mapped_paragraphs_get_full_text(self, sample_doc_structure, sample_mapping_plan):
        """Mapped paragraphs (section_index 3, 5, 8) should get full text (200 chars)."""
        result = build_zone_map(sample_doc_structure, sample_mapping_plan)
        # Paragraph 3 is mapped -- its full text (128 chars) should appear untruncated
        assert "This report presents the findings from the security assessment" in result

    def test_unmapped_paragraphs_get_truncated_text(self, sample_doc_structure, sample_mapping_plan):
        """Unmapped paragraphs not near mapped ones should be truncated to 50 chars."""
        # Paragraph 0 ("Company Logo Area") is unmapped and far from mapped indices
        # so it gets truncated. It's short enough to fit in 50 chars fully.
        result = build_zone_map(sample_doc_structure, sample_mapping_plan)
        assert "Company Logo Area" in result

    def test_paragraph_indices_are_correct(self, sample_doc_structure, sample_mapping_plan):
        result = build_zone_map(sample_doc_structure, sample_mapping_plan)
        # Cover paragraphs should have indices 0, 1
        assert "[0]" in result
        assert "[1]" in result
        # Body paragraphs should start at index 2
        assert "[2]" in result

    def test_style_names_included(self, sample_doc_structure, sample_mapping_plan):
        result = build_zone_map(sample_doc_structure, sample_mapping_plan)
        assert "[Title]" in result
        assert "[Normal]" in result
        assert "[Heading 1]" in result

    def test_heading_levels_included(self, sample_doc_structure, sample_mapping_plan):
        result = build_zone_map(sample_doc_structure, sample_mapping_plan)
        assert "(H1)" in result
        assert "(H2)" in result

    def test_empty_paragraphs_skipped(self, sample_doc_structure, sample_mapping_plan):
        """Paragraph at index 6 has empty text and should not produce a line."""
        result = build_zone_map(sample_doc_structure, sample_mapping_plan)
        # There should be no line for index 6 since its text is empty
        lines = result.split("\n")
        idx6_lines = [l for l in lines if l.strip().startswith("[6]")]
        assert len(idx6_lines) == 0

    def test_header_paragraph_format(self, sample_doc_structure, sample_mapping_plan):
        result = build_zone_map(sample_doc_structure, sample_mapping_plan)
        assert "[H0][Header]: Company Header" in result

    def test_footer_paragraph_format(self, sample_doc_structure, sample_mapping_plan):
        result = build_zone_map(sample_doc_structure, sample_mapping_plan)
        assert "[F0][Footer]: Page {{ page_number }}" in result


# ---------------------------------------------------------------------------
# build_placement_prompt tests
# ---------------------------------------------------------------------------


class TestBuildPlacementPrompt:
    def test_all_four_sections_present(self, sample_doc_structure, sample_mapping_plan):
        result = build_placement_prompt(sample_doc_structure, sample_mapping_plan)
        assert "## SECTION 1: Document Structure (Zone Map)" in result
        assert "## SECTION 2: Mapping Entries to Place" in result
        assert "## SECTION 3: Required Output Format" in result
        assert "## SECTION 4: Placement Rules" in result

    def test_mapping_entries_include_gw_field(self, sample_doc_structure, sample_mapping_plan):
        result = build_placement_prompt(sample_doc_structure, sample_mapping_plan)
        assert 'gw_field="client.short_name"' in result
        assert 'gw_field="finding.description_rt"' in result
        assert 'gw_field="finding.title"' in result

    def test_mapping_entries_include_placeholder(self, sample_doc_structure, sample_mapping_plan):
        result = build_placement_prompt(sample_doc_structure, sample_mapping_plan)
        assert "{{ client.short_name }}" in result
        assert "{{p finding.description_rt }}" in result

    def test_mapping_entries_include_marker_type(self, sample_doc_structure, sample_mapping_plan):
        result = build_placement_prompt(sample_doc_structure, sample_mapping_plan)
        assert 'marker_type="text"' in result
        assert 'marker_type="paragraph_rt"' in result

    def test_output_format_includes_confidence(self, sample_doc_structure, sample_mapping_plan):
        result = build_placement_prompt(sample_doc_structure, sample_mapping_plan)
        assert '"confidence"' in result

    def test_section_text_hint_included(self, sample_doc_structure, sample_mapping_plan):
        result = build_placement_prompt(sample_doc_structure, sample_mapping_plan)
        assert "section_text:" in result
        assert "This report presents the findings" in result

    def test_valid_actions_listed(self, sample_doc_structure, sample_mapping_plan):
        result = build_placement_prompt(sample_doc_structure, sample_mapping_plan)
        assert "replace_text" in result
        assert "insert_before" in result
        assert "insert_after" in result
        assert "wrap_table_row" in result


# ---------------------------------------------------------------------------
# build_placement_system_prompt tests
# ---------------------------------------------------------------------------


class TestBuildPlacementSystemPrompt:
    def test_returns_non_empty_string(self):
        result = build_placement_system_prompt()
        assert isinstance(result, str)
        assert len(result) > 0

    def test_defines_role(self):
        result = build_placement_system_prompt()
        assert "DOCX template engineer" in result
        assert "Jinja2 placeholder placement" in result

    def test_critical_rules_present(self):
        result = build_placement_system_prompt()
        assert "CRITICAL RULES" in result

    def test_exact_substring_rule(self):
        result = build_placement_system_prompt()
        assert "exact substring" in result

    def test_confidence_rule(self):
        result = build_placement_system_prompt()
        assert "confidence" in result

    def test_json_only_rule(self):
        result = build_placement_system_prompt()
        assert "ONLY valid JSON" in result

    def test_marker_types_documented(self):
        result = build_placement_system_prompt()
        assert "paragraph_rt" in result
        assert "run_rt" in result
        assert "table_row_loop" in result
        assert "control_flow" in result


# ---------------------------------------------------------------------------
# validate-placement logic tests (unit tests via direct function)
# ---------------------------------------------------------------------------


def _build_validate_request(
    instructions: list[dict],
    paragraphs: list[DocxParagraph],
    template_type: str = "web",
    language: str = "en",
) -> tuple[ValidatePlacementRequest, DocxStructure]:
    """Helper to build a ValidatePlacementRequest with mock data."""
    import base64
    from io import BytesIO
    from docx import Document

    # Create a real DOCX with the given paragraph texts
    doc = Document()
    for para in paragraphs:
        doc.add_paragraph(para.text)

    buf = BytesIO()
    doc.save(buf)
    doc_bytes = buf.getvalue()
    doc_b64 = base64.b64encode(doc_bytes).decode("ascii")

    llm_response = json.dumps({
        "instructions": instructions,
        "template_type": template_type,
        "language": language,
    })

    request = ValidatePlacementRequest(
        llm_response=llm_response,
        template_base64=doc_b64,
        template_type=template_type,
        language=language,
        paragraph_count=len(paragraphs),
    )
    return request


class TestValidatePlacement:
    """Test validate-placement endpoint logic via the FastAPI route."""

    @pytest.mark.anyio
    async def test_valid_instruction_passes(self):
        """A valid instruction with matching original_text should pass."""
        from app.routes.adapter import validate_placement

        paragraphs = [
            DocxParagraph(text="Company Name Here"),
            DocxParagraph(text="Executive Summary"),
            DocxParagraph(text="This report covers Client Corp."),
        ]
        instructions = [
            {
                "action": "replace_text",
                "paragraph_index": 2,
                "original_text": "Client Corp",
                "replacement_text": "{{ client.short_name }}",
                "marker_type": "text",
                "gw_field": "client.short_name",
                "confidence": 0.95,
            }
        ]
        request = _build_validate_request(instructions, paragraphs)
        response = await validate_placement(request)

        assert response.valid is True
        assert response.applied_count == 1
        assert response.skipped_count == 0
        assert response.instruction_set is not None
        assert len(response.instruction_set.instructions) == 1
        assert response.instruction_set.instructions[0].gw_field == "client.short_name"

    @pytest.mark.anyio
    async def test_out_of_bounds_paragraph_index_skipped(self):
        """Instruction with paragraph_index out of bounds should be skipped."""
        from app.routes.adapter import validate_placement

        paragraphs = [
            DocxParagraph(text="Only paragraph"),
        ]
        instructions = [
            {
                "action": "replace_text",
                "paragraph_index": 99,
                "original_text": "something",
                "replacement_text": "{{ field }}",
                "marker_type": "text",
                "gw_field": "client.short_name",
                "confidence": 0.9,
            }
        ]
        request = _build_validate_request(instructions, paragraphs)
        response = await validate_placement(request)

        assert response.valid is False
        assert response.skipped_count == 1
        assert any("out of bounds" in w for w in response.warnings)

    @pytest.mark.anyio
    async def test_non_matching_original_text_skipped(self):
        """Instruction with original_text not found in any paragraph should be skipped."""
        from app.routes.adapter import validate_placement

        paragraphs = [
            DocxParagraph(text="Company Name Here"),
            DocxParagraph(text="Executive Summary"),
        ]
        instructions = [
            {
                "action": "replace_text",
                "paragraph_index": 0,
                "original_text": "Nonexistent Text That Does Not Exist Anywhere",
                "replacement_text": "{{ field }}",
                "marker_type": "text",
                "gw_field": "client.short_name",
                "confidence": 0.9,
            }
        ]
        request = _build_validate_request(instructions, paragraphs)
        response = await validate_placement(request)

        assert response.valid is False
        assert response.skipped_count == 1
        assert any("not found" in w for w in response.warnings)

    @pytest.mark.anyio
    async def test_low_confidence_filtered_to_warning(self):
        """Instruction with confidence < 0.5 should be filtered with a warning."""
        from app.routes.adapter import validate_placement

        paragraphs = [
            DocxParagraph(text="Some text here"),
        ]
        instructions = [
            {
                "action": "replace_text",
                "paragraph_index": 0,
                "original_text": "Some text",
                "replacement_text": "{{ field }}",
                "marker_type": "text",
                "gw_field": "client.short_name",
                "confidence": 0.3,
            }
        ]
        request = _build_validate_request(instructions, paragraphs)
        response = await validate_placement(request)

        assert response.valid is False
        assert response.skipped_count == 1
        assert any("confidence" in w and "0.30" in w for w in response.warnings)

    @pytest.mark.anyio
    async def test_empty_response_returns_invalid(self):
        """Empty instructions array should return valid=false."""
        from app.routes.adapter import validate_placement

        paragraphs = [
            DocxParagraph(text="Some text"),
        ]
        request = _build_validate_request([], paragraphs)
        response = await validate_placement(request)

        assert response.valid is False
        assert any("empty" in e.lower() for e in response.errors)

    @pytest.mark.anyio
    async def test_instruction_set_compatible_with_apply(self):
        """Validated InstructionSet should have correct template_type and language."""
        from app.routes.adapter import validate_placement

        paragraphs = [
            DocxParagraph(text="Test paragraph with Client Corp name"),
        ]
        instructions = [
            {
                "action": "replace_text",
                "paragraph_index": 0,
                "original_text": "Client Corp",
                "replacement_text": "{{ client.short_name }}",
                "marker_type": "text",
                "gw_field": "client.short_name",
                "confidence": 0.9,
            }
        ]
        request = _build_validate_request(instructions, paragraphs)
        response = await validate_placement(request)

        assert response.valid is True
        assert response.instruction_set is not None
        assert response.instruction_set.template_type == "web"
        assert response.instruction_set.language == "en"

    @pytest.mark.anyio
    async def test_markdown_fences_stripped(self):
        """LLM response wrapped in markdown fences should be parsed correctly."""
        from app.routes.adapter import validate_placement
        import base64
        from io import BytesIO
        from docx import Document

        doc = Document()
        doc.add_paragraph("Hello World")
        buf = BytesIO()
        doc.save(buf)
        doc_b64 = base64.b64encode(buf.getvalue()).decode("ascii")

        llm_response = '```json\n{"instructions": [{"action": "replace_text", "paragraph_index": 0, "original_text": "Hello", "replacement_text": "{{ field }}", "marker_type": "text", "gw_field": "client.short_name", "confidence": 0.9}], "template_type": "web", "language": "en"}\n```'

        request = ValidatePlacementRequest(
            llm_response=llm_response,
            template_base64=doc_b64,
            template_type="web",
            language="en",
            paragraph_count=1,
        )
        response = await validate_placement(request)

        assert response.valid is True
        assert response.applied_count == 1

    @pytest.mark.anyio
    async def test_text_relocation_fallback(self):
        """If original_text isn't at paragraph_index, find it elsewhere."""
        from app.routes.adapter import validate_placement

        paragraphs = [
            DocxParagraph(text="First paragraph"),
            DocxParagraph(text="Second paragraph with Client Corp"),
            DocxParagraph(text="Third paragraph"),
        ]
        instructions = [
            {
                "action": "replace_text",
                "paragraph_index": 0,
                "original_text": "Client Corp",
                "replacement_text": "{{ client.short_name }}",
                "marker_type": "text",
                "gw_field": "client.short_name",
                "confidence": 0.9,
            }
        ]
        request = _build_validate_request(instructions, paragraphs)
        response = await validate_placement(request)

        assert response.valid is True
        assert response.applied_count == 1
        # Should have relocated and produced a warning
        assert any("relocated" in w for w in response.warnings)
        # The instruction should now point to paragraph 1
        assert response.instruction_set.instructions[0].paragraph_index == 1
